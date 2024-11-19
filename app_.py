import streamlit as st
from streamlit_drawable_canvas import st_canvas
from datetime import datetime, date, timedelta
import time
from PIL import Image as PILImage
import numpy as np
from docx import Document
from docx.shared import Inches
import smtplib
from email.message import EmailMessage
import shutil
import re
import os
from dotenv import load_dotenv
import traceback
# import io
import requests

st.set_page_config(
    page_title="Prevista - ESFA Form",
    page_icon="https://lirp.cdn-website.com/d8120025/dms3rep/multi/opt/social-image-88w.png",  # Path to logo
    # page_icon="📝",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# =========================================================================
# All Functions
# =========================================================================

# add render support along with st.secret
def get_secret(key):
    try:
        load_dotenv()
        # Attempt to get the secret from environment variables
        secret = os.environ.get(key)
        if secret is None:
            raise ValueError("Secret not found in environment variables")
        return secret
    except (ValueError, TypeError) as e:
        # If an error occurs, fall back to Streamlit secrets
        if hasattr(st, 'secrets'):
            return st.secrets.get(key)
        # If still not found, return None or handle as needed
        return None

# Sanitize the file name to avoid invalid characters
def sanitize_filename(filename):
    return re.sub(r'[<>:"/\\|?*]', '', filename)

def validate_inputs(inputs, mandatory_fields):
    """Check if all mandatory input fields are filled and return the list of missing fields."""
    missing_fields = []
    for key, value in inputs.items():
        if key in mandatory_fields and (value is None or value == '' or value == 0):
            missing_fields.append(key)
    return missing_fields


def is_valid_email(email):
    # Comprehensive regex for email validation
    pattern = r'''
        ^                         # Start of string
        (?!.*[._%+-]{2})          # No consecutive special characters
        [a-zA-Z0-9._%+-]{1,64}    # Local part: allowed characters and length limit
        (?<![._%+-])              # No special characters at the end of local part
        @                         # "@" symbol
        [a-zA-Z0-9.-]+            # Domain part: allowed characters
        (?<![.-])                 # No special characters at the end of domain
        \.[a-zA-Z]{2,}$           # Top-level domain with minimum 2 characters
    '''
    
    # Match the entire email against the pattern
    return re.match(pattern, email, re.VERBOSE) is not None

def replace_placeholders(template_file, modified_file, placeholder_values, resized_image_path_1, resized_image_path_2):
    try:
        print(f"Copying template file '{template_file}' to '{modified_file}'...")
        shutil.copy(template_file, modified_file)

        time.sleep(1)

        print(f"Opening document '{modified_file}'...")
        doc = Document(modified_file)

        # Function to convert value to string, handling datetime.date objects
        def convert_to_str(value):
            if isinstance(value, date):
                return value.strftime('%Y-%m-%d')  # Convert date to string
            return str(value)  # Convert other types to string

        # Compile regular expressions for all placeholders
        placeholders = {re.escape(key): convert_to_str(value) for key, value in placeholder_values.items()}
        placeholders_pattern = re.compile(r'\b(' + '|'.join(placeholders.keys()) + r')\b')

        # Replace placeholders in paragraphs
        print("Replacing placeholders in paragraphs...")
        for para in doc.paragraphs:
            original_text = para.text
            updated_text = placeholders_pattern.sub(lambda match: placeholders[re.escape(match.group(0))], para.text)
            if original_text != updated_text:
                print(f"Updated paragraph text: '{original_text}' -> '{updated_text}'")
                para.text = updated_text

        # Replace placeholders in tables
        print("Replacing placeholders in tables...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text
                        updated_text = placeholders_pattern.sub(lambda match: placeholders[re.escape(match.group(0))], para.text)
                        if original_text != updated_text:
                            print(f"Updated table cell text: '{original_text}' -> '{updated_text}'")
                            para.text = updated_text

                    # Inspect cell runs
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run_text = run.text
                            run_updated_text = placeholders_pattern.sub(lambda match: placeholders[re.escape(match.group(0))], run_text)
                            if run_text != run_updated_text:
                                print(f"Updated run text in table cell: '{run_text}' -> '{run_updated_text}'")
                                run.text = run_updated_text

        # Check and handle signature placeholders
        print("Inspecting document for 'p230' and 'p234' placeholders...")
        signature_placeholder_found = False

        # Function to insert an image when a placeholder is found
        def insert_signature_image(para, image_path):
            try:
                print(f"Adding picture to paragraph or cell from path: {image_path}")
                para.add_run().add_picture(image_path, width=Inches(2))
                print("Inserted signature image.")
                return True
            except Exception as img_e:
                print(f"An error occurred with image processing: {img_e}")
                return False

        # Check paragraphs for both 'p230' and 'p234'
        for para in doc.paragraphs:
            para_text = para.text.strip()  # Remove any extra spaces around text
            
            if 'p230' in para_text:
                print(f"Found 'p230' in paragraph: '{para_text}'")
                para.text = para_text.replace('p230', '').strip()
                signature_placeholder_found = insert_signature_image(para, resized_image_path_1)
            
            if 'p234' in para_text:
                print(f"Found 'p234' in paragraph: '{para_text}'")
                para.text = para_text.replace('p234', '').strip()
                signature_placeholder_found = insert_signature_image(para, resized_image_path_2)

        # Check table cells for 'p230' and 'p234'
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para_text = para.text.strip()
                        
                        if 'p230' in para_text:
                            print(f"Found 'p230' in table cell paragraph: '{para_text}'")
                            para.text = para_text.replace('p230', '').strip()
                            signature_placeholder_found = insert_signature_image(para, resized_image_path_1)
                        
                        if 'p234' in para_text:
                            print(f"Found 'p234' in table cell paragraph: '{para_text}'")
                            para.text = para_text.replace('p234', '').strip()
                            signature_placeholder_found = insert_signature_image(para, resized_image_path_2)

        if not signature_placeholder_found:
            print("No signature placeholder found.")

        # Save the modified document
        print(f"Saving modified document '{modified_file}'...")
        doc.save(modified_file)
        print(f"Document modification complete: '{modified_file}'")

    except Exception as e:
        print(f"An error occurred: {e}")


    # # file download button
    # with open(modified_file, 'rb') as f:
    #     file_contents = f.read()
    #     st.download_button(
    #         label="Download Your Response",
    #         data=file_contents,
    #         file_name=modified_file,
    #         mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    #     )


def resize_image_to_fit_cell(image, max_width, max_height):
    width, height = image.size
    aspect_ratio = width / height

    if width > max_width:
        width = max_width
        height = int(width / aspect_ratio)

    if height > max_height:
        height = max_height
        width = int(height * aspect_ratio)

    return image.resize((width, height))

# Function to send email with attachments (Handle Local + Uploaded)
def send_email_with_attachments(sender_email, sender_password, receiver_email, subject, body, files=None, local_file_path=None):
    msg = EmailMessage()
    msg['From'] = sender_email
    msg['To'] = receiver_email
    msg['Subject'] = subject
    msg.set_content(body, subtype='html')

    # Attach uploaded files
    if files:
        for uploaded_file in files:
            uploaded_file.seek(0)  # Move to the beginning of the UploadedFile
            msg.add_attachment(uploaded_file.read(), maintype='application', subtype='octet-stream', filename=uploaded_file.name)

    # Attach local file if specified
    if local_file_path:
        with open(local_file_path, 'rb') as f:
            file_data = f.read()
            file_name = local_file_path.split('/')[-1]
            msg.add_attachment(file_data, maintype='application', subtype='octet-stream', filename=file_name)

    # Use the SMTP server for sending the email
    with smtplib.SMTP('smtp.office365.com', 587) as server:
        server.starttls()
        server.login(sender_email, sender_password)
        server.send_message(msg)



# Function to add a checkbox with a file upload option
def add_checkbox_with_upload(label, key_prefix):
    checked = st.checkbox(label, key=f"{key_prefix}_checkbox")
    st.session_state.checkboxes[label] = checked
    if checked:
        st.text(f'Please upload a copy of your {label}')
        uploaded_file = st.file_uploader(f"Upload {label}", type=['pdf', 'jpg', 'jpeg', 'png', 'docx'], key=f"{key_prefix}_uploader")
        if uploaded_file is not None:
            file_identifier = (uploaded_file.name, uploaded_file.size)
            if file_identifier not in st.session_state.processed_files:
                st.session_state.files.append(uploaded_file)
                st.session_state.processed_files.add(file_identifier)
        # Second File Uploader
        uploaded_file_1 = st.file_uploader(f"Optional - Upload Back Side of The Document", type=['pdf', 'jpg', 'jpeg', 'png', 'docx'], key=f"{key_prefix}_uploader_1")
        if uploaded_file_1 is not None:
            file_identifier_1 = (uploaded_file_1.name, uploaded_file_1.size)
            if file_identifier_1 not in st.session_state.processed_files:
                st.session_state.files.append(uploaded_file_1)
                st.session_state.processed_files.add(file_identifier_1)
        return 'X'
    else:
        return '-'

# Function to handle file upload
def handle_file_upload(label):
    # global files
    st.text(f'Please upload a copy of your {label}')
    uploaded_file = st.file_uploader(f"Upload {label}", type=['pdf', 'jpg', 'jpeg', 'png', 'docx'])
    if uploaded_file is not None:
        st.session_state.files.append(uploaded_file)
        return 'X'
    else:
        return '-'
        
def calculate_age(born):
    today = date.today()
    return today.year - born.year - ((today.month, today.day) < (born.month, born.day))

def progress_bar(duration_seconds):
    """Displays a progress bar that fills over the specified duration."""
    progress_bar = st.progress(0)
    
    # Number of updates per second for smoother progress
    updates_per_second = 20
    # Time to wait between updates
    sleep_time = 1 / updates_per_second
    # Total number of updates
    total_updates = duration_seconds * updates_per_second
    
    for i in range(total_updates + 1):
        # Update the progress bar
        progress = i / total_updates
        progress_bar.progress(progress)
        # Sleep for the calculated time
        time.sleep(sleep_time)
    # st.write("Progress complete!")
# ==============================================================================================================================================

# Initialize session state

if 'files' not in st.session_state:
    st.session_state.files = []
if 'checkboxes' not in st.session_state:
    st.session_state.checkboxes = {}
if 'processed_files' not in st.session_state:
    st.session_state.processed_files = set()

if 'step' not in st.session_state: st.session_state.step = 7
if 'submission_done' not in st.session_state: st.session_state.submission_done = False
if 'unique_files' not in st.session_state: st.session_state.unique_files = []

# Step 2: Personal Information initialization
if 'title_mr' not in st.session_state: st.session_state.title_mr = ''
if 'title_mrs' not in st.session_state: st.session_state.title_mrs = ''
if 'title_miss' not in st.session_state: st.session_state.title_miss = ''
if 'title_ms' not in st.session_state: st.session_state.title_ms = ''
if 'title' not in st.session_state: st.session_state.title = 'Mr'
if 'first_name' not in st.session_state: st.session_state.first_name = ''
if 'middle_name' not in st.session_state: st.session_state.middle_name = ''
if 'family_name' not in st.session_state: st.session_state.family_name = ''
if 'learner_name' not in st.session_state: st.session_state.learner_name = ''
if 'qualification' not in st.session_state: st.session_state.qualification = 'High School Diploma'
if 'start_date' not in st.session_state: st.session_state.start_date = None
if 'end_date' not in st.session_state: st.session_state.end_date = None
if 'gender_m' not in st.session_state: st.session_state.gender_m = ''
if 'gender_f' not in st.session_state: st.session_state.gender_f = ''
if 'other_gender' not in st.session_state: st.session_state.other_gender = ''
if 'other_gender_text' not in st.session_state: st.session_state.other_gender_text = ''
if 'gender' not in st.session_state: st.session_state.gender = ''
if 'date_of_birth' not in st.session_state: st.session_state.date_of_birth = None
if 'current_age' not in st.session_state: st.session_state.current_age = ''
if 'current_age_text' not in st.session_state: st.session_state.current_age_text = ''

# Step 3: Ethnicity initialization
if 'ethnicity_category' not in st.session_state: st.session_state.ethnicity_category = 'White'  # Default value
if 'ethnicity' not in st.session_state: st.session_state.ethnicity = 'English/ Welsh/ Scottish/ N Irish/ British'  # Default ethnicity
if 'ethnicity_code' not in st.session_state: st.session_state.ethnicity_code = ''
if 'ethnicity_vars' not in st.session_state: st.session_state.ethnicity_vars = {f'ethnicity_{i}': '' for i in range(31, 49)}
if 'national_insurance_number' not in st.session_state: st.session_state.national_insurance_number = ''
if 'house_no_name_street' not in st.session_state: st.session_state.house_no_name_street = ''
if 'suburb_village' not in st.session_state: st.session_state.suburb_village = ''
if 'town_city' not in st.session_state: st.session_state.town_city = ''
if 'county' not in st.session_state: st.session_state.county = ''
if 'country_of_domicile' not in st.session_state: st.session_state.country_of_domicile = ''
if 'current_postcode' not in st.session_state: st.session_state.current_postcode = ''
if 'postcode_prior_enrollment' not in st.session_state: st.session_state.postcode_prior_enrollment = ''
if 'email_address' not in st.session_state: st.session_state.email_address = ''
if 'primary_telephone_number' not in st.session_state: st.session_state.primary_telephone_number = ''
if 'secondary_telephone_number' not in st.session_state: st.session_state.secondary_telephone_number = ''
if 'next_of_kin' not in st.session_state: st.session_state.next_of_kin = 'N/A'
if 'emergency_contact_phone_number' not in st.session_state: st.session_state.emergency_contact_phone_number = 'N/A'

# Step 4: Household initialization
if 'household_selections' not in st.session_state: st.session_state.household_selections = {}
if 'no_member_employed_with_children' not in st.session_state: st.session_state.no_member_employed_with_children = ''
if 'no_member_employed_without_children' not in st.session_state: st.session_state.no_member_employed_without_children = ''
if 'single_adult_household_with_children' not in st.session_state: st.session_state.single_adult_household_with_children = ''
if 'unemployed_single_adult_household' not in st.session_state: st.session_state.unemployed_single_adult_household = ''
if 'none_of_the_above' not in st.session_state: st.session_state.none_of_the_above = ''
if 'household_filled' not in st.session_state: st.session_state.household_filled = ''

# Step 5: LLDD, Health Problems, Other Disadvantaged Section
if 'has_disability' not in st.session_state: st.session_state.has_disability = ''
if 'no_disability' not in st.session_state: st.session_state.no_disability = ''
if 'vision_impairment_primary' not in st.session_state: st.session_state.vision_impairment_primary, st.session_state.vision_impairment_secondary, st.session_state.vision_impairment_tertiary = '', '', ''
if 'hearing_impairment_primary' not in st.session_state: st.session_state.hearing_impairment_primary, st.session_state.hearing_impairment_secondary, st.session_state.hearing_impairment_tertiary = '', '', ''
if 'mobility_impairment_primary' not in st.session_state: st.session_state.mobility_impairment_primary, st.session_state.mobility_impairment_secondary, st.session_state.mobility_impairment_tertiary = '', '', ''
if 'complex_disabilities_primary' not in st.session_state: st.session_state.complex_disabilities_primary, st.session_state.complex_disabilities_secondary, st.session_state.complex_disabilities_tertiary = '', '', ''
if 'social_emotional_difficulties_primary' not in st.session_state: st.session_state.social_emotional_difficulties_primary, st.session_state.social_emotional_difficulties_secondary, st.session_state.social_emotional_difficulties_tertiary = '', '', ''
if 'mental_health_difficulty_primary' not in st.session_state: st.session_state.mental_health_difficulty_primary, st.session_state.mental_health_difficulty_secondary, st.session_state.mental_health_difficulty_tertiary = '', '', ''
if 'moderate_learning_difficulty_primary' not in st.session_state: st.session_state.moderate_learning_difficulty_primary, st.session_state.moderate_learning_difficulty_secondary, st.session_state.moderate_learning_difficulty_tertiary = '', '', ''
if 'severe_learning_difficulty_primary' not in st.session_state: st.session_state.severe_learning_difficulty_primary, st.session_state.severe_learning_difficulty_secondary, st.session_state.severe_learning_difficulty_tertiary = '', '', ''
if 'dyslexia_primary' not in st.session_state: st.session_state.dyslexia_primary, st.session_state.dyslexia_secondary, st.session_state.dyslexia_tertiary = '', '', ''
if 'dyscalculia_primary' not in st.session_state: st.session_state.dyscalculia_primary, st.session_state.dyscalculia_secondary, st.session_state.dyscalculia_tertiary = '', '', ''
if 'autism_spectrum_primary' not in st.session_state: st.session_state.autism_spectrum_primary, st.session_state.autism_spectrum_secondary, st.session_state.autism_spectrum_tertiary = '', '', ''
if 'aspergers_primary' not in st.session_state: st.session_state.aspergers_primary, st.session_state.aspergers_secondary, st.session_state.aspergers_tertiary = '', '', ''
if 'temporary_disability_primary' not in st.session_state: st.session_state.temporary_disability_primary, st.session_state.temporary_disability_secondary, st.session_state.temporary_disability_tertiary = '', '', ''
if 'speech_communication_needs_primary' not in st.session_state: st.session_state.speech_communication_needs_primary, st.session_state.speech_communication_needs_secondary, st.session_state.speech_communication_needs_tertiary = '', '', ''
if 'physical_disability_primary' not in st.session_state: st.session_state.physical_disability_primary, st.session_state.physical_disability_secondary, st.session_state.physical_disability_tertiary = '', '', ''
if 'specific_learning_difficulty_primary' not in st.session_state: st.session_state.specific_learning_difficulty_primary, st.session_state.specific_learning_difficulty_secondary, st.session_state.specific_learning_difficulty_tertiary = '', '', ''
if 'medical_condition_primary' not in st.session_state: st.session_state.medical_condition_primary, st.session_state.medical_condition_secondary, st.session_state.medical_condition_tertiary = '', '', ''
if 'other_learning_difficulty_primary' not in st.session_state: st.session_state.other_learning_difficulty_primary, st.session_state.other_learning_difficulty_secondary, st.session_state.other_learning_difficulty_tertiary = '', '', ''
if 'other_disability_primary' not in st.session_state: st.session_state.other_disability_primary, st.session_state.other_disability_secondary, st.session_state.other_disability_tertiary = '', '', ''
if 'prefer_not_to_say' not in st.session_state: st.session_state.prefer_not_to_say = ''
if 'additional_info' not in st.session_state: st.session_state.additional_info = ''

# Step 6: Referral Source Section
if 'internally_sourced' not in st.session_state: st.session_state.internally_sourced = ''
if 'recommendation' not in st.session_state: st.session_state.recommendation = ''
if 'event' not in st.session_state: st.session_state.event = ''
if 'self_referral' not in st.session_state: st.session_state.self_referral = ''
if 'family_friends' not in st.session_state: st.session_state.family_friends = ''
if 'other' not in st.session_state: st.session_state.other = ''
if 'website' not in st.session_state: st.session_state.website = ''
if 'promotional_material' not in st.session_state: st.session_state.promotional_material = ''
if 'jobcentre_plus' not in st.session_state: st.session_state.jobcentre_plus = ''
if 'event_specify' not in st.session_state: st.session_state.event_specify = ''
if 'other_specify' not in st.session_state: st.session_state.other_specify = ''
if 'internally_sourced_val' not in st.session_state: st.session_state.internally_sourced_val = ''
if 'recommendation_val' not in st.session_state: st.session_state.recommendation_val = ''
if 'event_val' not in st.session_state: st.session_state.event_val = ''
if 'self_referral_val' not in st.session_state: st.session_state.self_referral_val = ''
if 'family_friends_val' not in st.session_state: st.session_state.family_friends_val = ''
if 'other_val' not in st.session_state: st.session_state.other_val = ''
if 'website_val' not in st.session_state: st.session_state.website_val = ''
if 'promotional_material_val' not in st.session_state: st.session_state.promotional_material_val = ''
if 'jobcentre_plus_val' not in st.session_state: st.session_state.jobcentre_plus_val = ''
if 'specify_referral' not in st.session_state: st.session_state.specify_referral = ''

# Step 7: Employment and Monitoring Information Section
if 'benefit_claim_date_val' not in st.session_state: st.session_state.benefit_claim_date_val = None
if 'resident_y' not in st.session_state: st.session_state.resident_y = ''
if 'resident_n' not in st.session_state: st.session_state.resident_n = ''
if 'country_of_birth' not in st.session_state: st.session_state.country_of_birth = ''
if 'years_in_uk' not in st.session_state: st.session_state.years_in_uk = 0
if 'unemployed_val' not in st.session_state: st.session_state.unemployed_val, st.session_state.economically_inactive_val, st.session_state.employed_val = '', '', ''
if 'up_to_12_months_val' not in st.session_state: st.session_state.up_to_12_months_val, st.session_state.twelve_months_or_longer_val = '', ''
if 'jcp_dwp_val' not in st.session_state: st.session_state.jcp_dwp_val, st.session_state.careers_service_val, st.session_state.third_party_val, st.session_state.other_evidence_val = '', '', '', ''
if 'job_position' not in st.session_state: st.session_state.job_position = ''
if 'job_start_date' not in st.session_state: st.session_state.job_start_date = None
if 'tp_name' not in st.session_state: st.session_state.tp_name = ''
if 'inactive_status_val' not in st.session_state: st.session_state.inactive_status_val, st.session_state.inactive_evidence_type_val, st.session_state.inactive_evidence_date_val = 'N', '', None
if 'employer_name_val' not in st.session_state: st.session_state.employer_name_val, st.session_state.employer_address_1_val, st.session_state.employer_address_2_val = '', '', ''
if 'employer_address_3_val' not in st.session_state: st.session_state.employer_address_3_val, st.session_state.employer_postcode_val, st.session_state.employer_contact_name_val = '', '', ''
if 'employer_contact_position_val' not in st.session_state: st.session_state.employer_contact_position_val, st.session_state.employer_contact_email_val, st.session_state.employer_contact_phone_val = '', '', ''
if 'employer_edrs_number_val' not in st.session_state: st.session_state.employer_edrs_number_val, st.session_state.living_wage_val, st.session_state.employment_hours_val_0, st.session_state.employment_hours_val_6 = '', '', '', ''
if 'claiming_benefits_val' not in st.session_state: st.session_state.claiming_benefits_val, st.session_state.sole_claimant_val, st.session_state.benefits_list_val = '', '', ''
if 'other_benefit_val' not in st.session_state: st.session_state.other_benefit_val = ''
if 'universal_credit_val' not in st.session_state: st.session_state.universal_credit_val = ''
if 'job_seekers_allowance_val' not in st.session_state: st.session_state.job_seekers_allowance_val = ''
if 'employment_support_allowance_val' not in st.session_state: st.session_state.employment_support_allowance_val = ''
if 'incapacity_benefit_val' not in st.session_state: st.session_state.incapacity_benefit_val = ''
if 'personal_independence_payment_val' not in st.session_state: st.session_state.personal_independence_payment_val = ''
if 'nationality' not in st.session_state: st.session_state.nationality = ''
if 'hold_settled_status' not in st.session_state: st.session_state.hold_settled_status, st.session_state.hold_pre_settled_status, st.session_state.hold_leave_to_remain = '', '', ''
if 'not_nationality' not in st.session_state: st.session_state.not_nationality, st.session_state.passport_non_eu, st.session_state.letter_uk_immigration, st.session_state.passport_endorsed, st.session_state.identity_card, st.session_state.country_of_issue, st.session_state.id_document_reference_number, st.session_state.e01_date_of_issue, st.session_state.e01_date_of_expiry, st.session_state.e01_additional_notes = '', '', '', '', '', '', '', '', '', ''
if 'full_uk_passport' not in st.session_state: st.session_state.full_uk_passport, st.session_state.full_eu_passport, st.session_state.national_identity_card = '', '', ''
if 'full_passport_eu' not in st.session_state: st.session_state.full_passport_eu = ''
if 'national_id_card_eu' not in st.session_state: st.session_state.national_id_card_eu = ''
if 'firearms_certificate' not in st.session_state: st.session_state.firearms_certificate = ''
if 'birth_adoption_certificate' not in st.session_state: st.session_state.birth_adoption_certificate = ''
if 'e02_drivers_license' not in st.session_state: st.session_state.e02_drivers_license = ''
if 'edu_institution_letter' not in st.session_state: st.session_state.edu_institution_letter = ''
if 'e02_employment_contract' not in st.session_state: st.session_state.e02_employment_contract = ''
if 'state_benefits_letter' not in st.session_state: st.session_state.state_benefits_letter = ''
if 'pension_statement' not in st.session_state: st.session_state.pension_statement = ''
if 'northern_ireland_voters_card' not in st.session_state: st.session_state.northern_ireland_voters_card = ''
if 'e02_other_evidence_text' not in st.session_state: st.session_state.e02_other_evidence_text = ''
if 'e03_drivers_license' not in st.session_state: st.session_state.e03_drivers_license = ''
if 'bank_statement' not in st.session_state: st.session_state.bank_statement = ''
if 'e03_pension_statement' not in st.session_state: st.session_state.e03_pension_statement = ''
if 'mortgage_statement' not in st.session_state: st.session_state.mortgage_statement = ''
if 'utility_bill' not in st.session_state: st.session_state.utility_bill = ''
if 'council_tax_statement' not in st.session_state: st.session_state.council_tax_statement = ''
if 'electoral_role_evidence' not in st.session_state: st.session_state.electoral_role_evidence = ''
if 'homeowner_letter' not in st.session_state: st.session_state.homeowner_letter = ''
if 'e02_date_of_issue' not in st.session_state: st.session_state.e02_date_of_issue = None
if 'e03_date_of_issue' not in st.session_state: st.session_state.e03_date_of_issue = None
if 'e03_other_evidence_text' not in st.session_state: st.session_state.e03_other_evidence_text = ''
if 'latest_payslip' not in st.session_state: st.session_state.latest_payslip = '-'
if 'e04_employment_contract' not in st.session_state: st.session_state.e04_employment_contract = '-'
if 'confirmation_from_employer' not in st.session_state: st.session_state.confirmation_from_employer = '-'
if 'redundancy_notice' not in st.session_state: st.session_state.redundancy_notice = '-'
if 'sa302_declaration' not in st.session_state: st.session_state.sa302_declaration = '-'
if 'ni_contributions' not in st.session_state: st.session_state.ni_contributions = '-'
if 'business_records' not in st.session_state: st.session_state.business_records = '-'
if 'companies_house_records' not in st.session_state: st.session_state.companies_house_records = '-'
if 'other_evidence_employed' not in st.session_state: st.session_state.other_evidence_employed = '-'
if 'unemployed' not in st.session_state: st.session_state.unemployed = '-'
if 'e04_date_of_issue' not in st.session_state: st.session_state.e04_date_of_issue = ''
if 'qualification_or_training_y' not in st.session_state: st.session_state.qualification_or_training_y, st.session_state.qualification_or_training_n = 'Y', ''
if 'course_details' not in st.session_state: st.session_state.course_details = ''
if 'funding_details' not in st.session_state: st.session_state.funding_details = ''
if 'p58' not in st.session_state: st.session_state.p58 = '-'
if 'p59' not in st.session_state: st.session_state.p59 = '-'
if 'p60' not in st.session_state: st.session_state.p60 = '-'
if 'p60z' not in st.session_state: st.session_state.p60z = '-'
if 'p60a' not in st.session_state: st.session_state.p60a = '-'
if 'p61' not in st.session_state: st.session_state.p61 = '-'
if 'p61z' not in st.session_state: st.session_state.p61z = '-'
if 'p61a' not in st.session_state: st.session_state.p61a = '-'
if 'p62' not in st.session_state: st.session_state.p62 = '-'
if 'p63' not in st.session_state: st.session_state.p63 = '-'
if 'p63z' not in st.session_state: st.session_state.p63z = '-'
if 'p63a' not in st.session_state: st.session_state.p63a = '-'
if 'p63b' not in st.session_state: st.session_state.p63b = '-'
if 'p64' not in st.session_state: st.session_state.p64 = '-'
if 'e02_filled' not in st.session_state: st.session_state.e02_filled = ''
if 'e03_drivers_license' not in st.session_state: st.session_state.e03_drivers_license = ''
if 'bank_statement' not in st.session_state: st.session_state.bank_statement = ''
if 'e03_pension_statement' not in st.session_state: st.session_state.e03_pension_statement = ''
if 'mortgage_statement' not in st.session_state: st.session_state.mortgage_statement = ''
if 'utility_bill' not in st.session_state: st.session_state.utility_bill = ''
if 'council_tax_statement' not in st.session_state: st.session_state.council_tax_statement = ''
if 'electoral_role_evidence' not in st.session_state: st.session_state.electoral_role_evidence = ''
if 'homeowner_letter' not in st.session_state: st.session_state.homeowner_letter = ''
if 'e03_filled' not in st.session_state: st.session_state.e03_filled = ''
if 'passport_non_eu_checked' not in st.session_state: st.session_state.passport_non_eu_checked = ''


# Step 9: Current Skills, Experience, and IAG
if 'p93' not in st.session_state: st.session_state.p93, st.session_state.p94, st.session_state.p95, st.session_state.p96, st.session_state.p97, st.session_state.p98 = '-', '-', '-', '-', '-', '-'
if 'job_role_activities' not in st.session_state: st.session_state.job_role_activities = 'No job.'
if 'career_aspirations' not in st.session_state: st.session_state.career_aspirations = ''
if 'training_qualifications_needed' not in st.session_state: st.session_state.training_qualifications_needed = '    '
if 'barriers_to_achieving_aspirations' not in st.session_state: st.session_state.barriers_to_achieving_aspirations = '    '
if 'selected_levels' not in st.session_state: st.session_state.selected_levels = ''
if 'current_job' not in st.session_state: st.session_state.current_job = ''

# Step 10: Privacy Notice
# Initialize variables for contact preferences
if 'contact_surveys_val' not in st.session_state: st.session_state.contact_surveys_val = ''
if 'contact_phone_val' not in st.session_state: st.session_state.contact_phone_val = ''
if 'contact_email_val' not in st.session_state: st.session_state.contact_email_val = ''
if 'contact_post_val' not in st.session_state: st.session_state.contact_post_val = ''

# mandatory fields validation
# exclude_fields = {}     
# mandatory_fields = []


# Define a function to calculate progress and percentage
def get_progress(step, total_steps):
    return int((step / total_steps) * 100)
# Define the total number of steps
total_steps = 12
# Calculate the current progress
progress = get_progress(st.session_state.step, total_steps)
# Display the progress bar and percentage
st.write(f"Progress: {progress}%")
st.progress(progress)

if st.session_state.step == 1:
    st.image('header/header.jpg', use_column_width=True)

    st.title('Welcome')

    # Initialize selected_option if it doesn't exist
    if 'selected_option' not in st.session_state:
        st.session_state.selected_option = "    "  # Default value

    # Add question with a dropdown menu
    support_options = [
        "    ", "Family Ties", "Catalyst", "Futures", "Innovators", "Alphabets", "Winners", 
        "Ealing Job Centre", "Ealing Council", "Brent Council", 
        "Brent JCP", "Tower Hamlets JCP", "Tower Hamlets Council", 
        "Oxfordshire JCP", "Surrey JCPs"
    ]

    # Use the selected_option from the session state to set the default value
    st.session_state.selected_option = st.selectbox(
        "Who is supporting you to fill this form?", 
        support_options, 
        index=support_options.index(st.session_state.selected_option)  # Set default value based on session state
    )

    st.subheader('Please fill out the complete form')
    st.text('Please click Next to begin.')

    if st.button("Next"):
        if st.session_state.selected_option != '    ':
            st.session_state.step = 2
            st.experimental_rerun()
        else:
            st.warning("Please Choose a Valid Support Option.")

elif st.session_state.step == 2:
    st.title("> 1: Personal Information")
    
    st.session_state.title = st.radio(
        "Title",
        ["Mr", "Mrs", "Miss", "Ms"],
        index=["Mr", "Mrs", "Miss", "Ms"].index(st.session_state.title)  # Set default index based on session state
    )
    if st.session_state.title == "Mr":
        st.session_state.title_mr = 'X'
    elif st.session_state.title == "Mrs":
        st.session_state.title_mrs = 'X'
    elif st.session_state.title == "Miss":
        st.session_state.title_miss = 'X'
    elif st.session_state.title == "Ms":
        st.session_state.title_ms = 'X'

    # Input fields for names
    st.session_state.first_name = st.text_input('First Name', value=st.session_state.first_name)
    st.session_state.middle_name = st.text_input('Middle Name (optional)', value=st.session_state.middle_name)
    st.session_state.family_name = st.text_input('Family Name', value=st.session_state.family_name)

    # Construct learner name
    st.session_state.learner_name = f"{st.session_state.first_name} {st.session_state.middle_name} {st.session_state.family_name}".strip()

    # Check if dob is a string and convert it back to a date object
    if isinstance(st.session_state.get("start_date"), str):
        st.session_state.start_date = datetime.strptime(st.session_state.get("start_date"), "%d-%m-%Y").date()

    st.session_state.start_date = st.date_input(
        label="Aim Start Date",  # Label for the field
        value=st.session_state.get("start_date"),  # Correctly access from session state
        min_value=date(1900, 1, 1),  # Minimum selectable date
        max_value=date.today(),  # Maximum selectable date
        help="Choose a date",  # Tooltip text
        format='DD/MM/YYYY'
    )

    if isinstance(st.session_state.get("end_date"), str):
        st.session_state.end_date = datetime.strptime(st.session_state.get("end_date"), "%d-%m-%Y").date()

    st.session_state.end_date = st.date_input(
        label="Expected Aim End Date",  # Label for the field
        value=st.session_state.get("end_date"),  # Correctly access from session state
        min_value=date(1900, 1, 1),  # Minimum selectable date
        max_value=date.today(),  # Maximum selectable date
        help="Choose a date",  # Tooltip text
        format='DD/MM/YYYY'
    )

    st.session_state.qualification = st.selectbox('Qualification', [
        'High School Diploma', 'Bachelor\'s Degree', 'Master\'s Degree', 'PhD', 'Other'
    ], index=['High School Diploma', 'Bachelor\'s Degree', 'Master\'s Degree', 'PhD', 'Other'].index(st.session_state.qualification))

    # Initialize gender variables
    st.session_state.gender_m, st.session_state.gender_f, st.session_state.other_gender = '', '', ''
    
    # Radio button for gender selection
    st.session_state.gender = st.radio(
        "Gender", ["M", "F", "Other"],
        index=["M", "F", "Other"].index(st.session_state.gender) if st.session_state.gender else 0  # Set default index
    )

    # Conditional input for "Other" gender option
    if st.session_state.gender == "M":
        st.session_state.gender_m = 'M'
        st.session_state.gender_f = ''
        st.session_state.other_gender = ''
    elif st.session_state.gender == "F":
        st.session_state.gender_f = 'F'
        st.session_state.gender_m = ''
        st.session_state.other_gender = ''
    elif st.session_state.gender == "Other":
        st.session_state.other_gender = 'Other'
        st.session_state.other_gender_text = st.text_input("If Other, please state", value=st.session_state.other_gender_text)


    if isinstance(st.session_state.get("date_of_birth"), str):
        st.session_state.date_of_birth = datetime.strptime(st.session_state.get("date_of_birth"), "%d-%m-%Y").date()

    st.session_state.date_of_birth = st.date_input(
        label="Date of Birth",
        value=st.session_state.get("date_of_birth"),
        min_value=date(1900, 1, 1),
        max_value=date.today(),
        help="Choose a date",
        format='DD/MM/YYYY'
    )
    
    # Next and Back buttons for navigation
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        if (st.session_state.first_name and st.session_state.family_name):

            # Convert the selected date to the desired string format (DD-MM-YYYY) only when proceeding to the next step
            st.session_state.start_date = st.session_state.start_date.strftime("%d-%m-%Y")
            st.session_state.end_date = st.session_state.end_date.strftime("%d-%m-%Y")
            st.session_state.current_age = calculate_age(st.session_state.date_of_birth)
            st.session_state.date_of_birth = st.session_state.date_of_birth.strftime("%d-%m-%Y")
            
            st.session_state.step = 3
            st.experimental_rerun()
        else:
            st.warning("Please fill in all fields before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 1  # Go back to the previous step (Section 1)
        st.experimental_rerun()


elif st.session_state.step == 3:

    st.title("> 2: Ethnicity")
    ethnicity_options = {
        'White': {
            'English/ Welsh/ Scottish/ N Irish/ British': '31',
            'Irish': '32',
            'Roma, Gypsy or Irish Traveller': '33',
            'Any other white background': '34'
        },
        'Mixed/ Multiple ethnic group': {
            'White and Black Caribbean': '35',
            'White and Black African': '36',
            'White and Asian': '37',
            'Any other mixed/ multiple ethnic background': '38'
        },
        'Asian/ Asian British': {
            'Bangladeshi': '41',
            'Chinese': '42',
            'Indian': '39',
            'Pakistani': '40',
            'Any other Asian background': '43'
        },
        'Black/ African/ Caribbean/ Black British': {
            'African': '44',
            'Caribbean': '45',
            'Any Other Black/ African/ Caribbean background': '46'
        },
        'Other Ethnic Group': {
            'Arab': '47',
            'Any other ethnic group': '48'
        }
    }

    # Select ethnicity category and ethnicity
    if 'ethnicity_category' not in st.session_state or st.session_state.ethnicity_category not in ethnicity_options:
        st.session_state.ethnicity_category = list(ethnicity_options.keys())[0]  # Default value

    st.session_state.ethnicity_category = st.selectbox(
        'Select Ethnicity Category',
        list(ethnicity_options.keys()),
        index=list(ethnicity_options.keys()).index(st.session_state.ethnicity_category)  # Set default index
    )

    # Ensure ethnicity is retained
    if 'ethnicity' not in st.session_state or st.session_state.ethnicity not in ethnicity_options[st.session_state.ethnicity_category]:
        st.session_state.ethnicity = list(ethnicity_options[st.session_state.ethnicity_category].keys())[0]  # Default to first option

    st.session_state.ethnicity = st.selectbox(
        'Select Ethnicity',
        list(ethnicity_options[st.session_state.ethnicity_category].keys()),
        index=list(ethnicity_options[st.session_state.ethnicity_category].keys()).index(st.session_state.ethnicity)  # Set default index
    )

    # Retrieve and convert ethnicity code to integer
    ethnicity_code_str = ethnicity_options[st.session_state.ethnicity_category][st.session_state.ethnicity]
    st.session_state.ethnicity_code = int(ethnicity_code_str)  # Ensure it is an integer
    st.write(f'Ethnicity Code: {st.session_state.ethnicity_code}')

    st.session_state.ethnicity_vars = {f'ethnicity_{i}': '' for i in range(31, 49)}

    # Set the corresponding ethnicity variable to 'X'
    if st.session_state.ethnicity_code in range(31, 49):
        st.session_state.ethnicity_vars[f'ethnicity_{st.session_state.ethnicity_code}'] = 'X'

    st.session_state.national_insurance_number = st.text_input("National Insurance Number", value=st.session_state.national_insurance_number)

    # Text inputs for address and contact details
    st.session_state.house_no_name_street = st.text_input("House No./Name & Street", value=st.session_state.house_no_name_street)
    st.session_state.suburb_village = st.text_input("Suburb / Village (Optional)", value=st.session_state.suburb_village)
    st.session_state.town_city = st.text_input("Town / City", value=st.session_state.town_city)
    st.session_state.county = st.text_input("County (optional)", value=st.session_state.county)
    st.session_state.country_of_domicile = st.text_input("Country of Domicile", value=st.session_state.country_of_domicile)
    st.session_state.current_postcode = st.text_input("Current Postcode", value=st.session_state.current_postcode)
    st.session_state.postcode_prior_enrollment = st.text_input("Postcode Prior to Enrolment", value=st.session_state.postcode_prior_enrollment)
    st.session_state.email_address = st.text_input("Email Address", value=st.session_state.email_address).strip().replace(" ", "_").lower()
    st.session_state.primary_telephone_number = st.text_input("Primary Telephone Number", value=st.session_state.primary_telephone_number)
    st.session_state.secondary_telephone_number = st.text_input("Secondary Telephone Number (optional)", value=st.session_state.secondary_telephone_number)
    st.session_state.next_of_kin = st.text_input("Next of kin/Emergency contact", value=st.session_state.next_of_kin)
    st.session_state.emergency_contact_phone_number = st.text_input("Emergency Contact Phone Number", value=st.session_state.emergency_contact_phone_number)

    # Next and Back buttons for navigation
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        if (is_valid_email(st.session_state.email_address)):
            if (st.session_state.national_insurance_number and
                st.session_state.house_no_name_street and
                st.session_state.town_city and
                st.session_state.country_of_domicile and
                st.session_state.current_postcode and
                st.session_state.postcode_prior_enrollment and
                st.session_state.primary_telephone_number):
                st.session_state.step = 4
                st.experimental_rerun()
            else:
                st.warning("Please fill in all fields before proceeding.")
        else:
            st.warning("Please enter valid email address.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 2  # Go back to the previous step (Section 2)
        st.experimental_rerun()


elif st.session_state.step == 4:
    # Household Situation Section
    st.title("> 3: Household")
    st.header('Household Situation')
    st.subheader('Please select the most relevant options. (Tick ALL relevant boxes)')

    household_options = {
        '1 - No household member in employment with one or more dependent children': 'JH, JH+DC',
        '2 - No household member in employment with no dependent children': 'JH',
        '3 - Participant lives in a single adult household with dependent children': 'SAH+DC',
        '4 - Learner lives in single unemployed adult household with dependent children': 'JH, SAH+DC',
        '99 - None of the above apply': 'N/A'
    }

    # Initialize household selections if not present in session state
    if 'household_selections' not in st.session_state:
        st.session_state.household_selections = {option: False for option in household_options.keys()}

    # Store household selections
    for option in household_options.keys():
        st.session_state.household_selections[option] = st.checkbox(option, value=st.session_state.household_selections[option], key=household_options[option])

    # Initialize relevant variables with empty string values
    st.session_state.no_member_employed_with_children = ''
    st.session_state.no_member_employed_without_children = ''
    st.session_state.single_adult_household_with_children = ''
    st.session_state.unemployed_single_adult_household = ''
    st.session_state.none_of_the_above = ''

    # Set variables based on selections
    if st.session_state.household_selections.get('1 - No household member in employment with one or more dependent children'):
        st.session_state.no_member_employed_with_children = 'X'
    if st.session_state.household_selections.get('2 - No household member in employment with no dependent children'):
        st.session_state.no_member_employed_without_children = 'X'
    if st.session_state.household_selections.get('3 - Participant lives in a single adult household with dependent children'):
        st.session_state.single_adult_household_with_children = 'X'
    if st.session_state.household_selections.get('4 - Learner lives in single unemployed adult household with dependent children'):
        st.session_state.unemployed_single_adult_household = 'X'
    if st.session_state.household_selections.get('99 - None of the above apply'):
        st.session_state.none_of_the_above = 'X'
        
    # Check if at least one checkbox is selected
    if any(st.session_state.household_selections.values()):
        st.session_state.household_filled = 'filled'
    else:
        st.session_state.household_filled = None

    # Next and Back buttons for navigation
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        if (st.session_state.household_filled):
            st.session_state.step = 5
            st.experimental_rerun()
        else:
            st.warning("Please tick at lease one box before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 3  # Go back to the previous step (Section 2)
        st.experimental_rerun()
        
elif st.session_state.step == 5:
    st.title("> 4: LLDD, Health Problems, Other Disadvantaged Section")

    # LLDD, Health Problems, Other Disadvantaged Section
    st.header('LLDD, Health Problems, Other Disadvantaged')

    # Long term disability, health problem, or learning difficulties
    st.write('Do you consider yourself to have a long term disability, health problem or any learning difficulties? Choose the correct option. If Yes enter code in Primary LLDD or HP; you can add multiple LLDD or HP but primary must be recorded if Yes selected.')
    
    # Initialize disability variable if it doesn't exist
    if 'disability' not in st.session_state:
        st.session_state.disability = 'N'  # Default value

    st.session_state.disability = st.radio('Choose the correct option:', ['N', 'Y'], index=['N', 'Y'].index(st.session_state.disability))

    # Initialize variables for disability options
    st.session_state.has_disability, st.session_state.no_disability = '', ''

    # Set variables based on user selection
    if st.session_state.disability == 'Y':
        st.session_state.has_disability, st.session_state.no_disability = 'Y', ''

        # LLDD or Health Problem Types
        st.subheader('LLDD or Health Problem Type')

        # Health problem types data
        data = [
            ('Vision impairment (4)', 'vision_primary', 'vision_secondary', 'vision_tertiary'),
            ('Hearing impairment (5)', 'hearing_primary', 'hearing_secondary', 'hearing_tertiary'),
            ('Disability affecting mobility (6)', 'mobility_primary', 'mobility_secondary', 'mobility_tertiary'),
            ('Profound complex disabilities (7)', 'complex_primary', 'complex_secondary', 'complex_tertiary'),
            ('Social and emotional difficulties (8)', 'social_primary', 'social_secondary', 'social_tertiary'),
            ('Mental health difficulty (9)', 'mental_primary', 'mental_secondary', 'mental_tertiary'),
            ('Moderate learning difficulty (10)', 'moderate_primary', 'moderate_secondary', 'moderate_tertiary'),
            ('Severe learning difficulty (11)', 'severe_primary', 'severe_secondary', 'severe_tertiary'),
            ('Dyslexia (12)', 'dyslexia_primary', 'dyslexia_secondary', 'dyslexia_tertiary'),
            ('Dyscalculia (13)', 'dyscalculia_primary', 'dyscalculia_secondary', 'dyscalculia_tertiary'),
            ('Autism spectrum disorder (14)', 'autism_primary', 'autism_secondary', 'autism_tertiary'),
            ('Asperger\'s syndrome (15)', 'aspergers_primary', 'aspergers_secondary', 'aspergers_tertiary'),
            ('Temporary disability after illness (for example post-viral) or accident (16)', 'temporary_primary', 'temporary_secondary', 'temporary_tertiary'),
            ('Speech, Language and Communication Needs (17)', 'speech_primary', 'speech_secondary', 'speech_tertiary'),
            ('Other physical disability (18)', 'physical_primary', 'physical_secondary', 'physical_tertiary'),
            ('Other specific learning difficulty (e.g. Dyspraxia) (19)', 'specific_primary', 'specific_secondary', 'specific_tertiary'),
            ('Other medical condition (for example epilepsy, asthma, diabetes) (20)', 'medical_primary', 'medical_secondary', 'medical_tertiary'),
            ('Other learning difficulty (90)', 'other_learning_primary', 'other_learning_secondary', 'other_learning_tertiary'),
            ('Other disability (97)', 'other_disability_primary', 'other_disability_secondary', 'other_disability_tertiary'),
            ('Prefer not to say (98)', 'prefer_not_to_say', '', '')
        ]

        # Create checkboxes and map them to variables explicitly
        for label, primary, secondary, tertiary in data:
            st.write(f'**{label}**')
            
            # Create checkboxes with unique keys and check current selections
            primary_checked = st.checkbox('Primary', key=f'{primary}_primary_checkbox', value='X' if st.session_state.get(primary) == 'X' else False)
            secondary_checked = st.checkbox('Secondary', key=f'{secondary}_secondary_checkbox', value='X' if st.session_state.get(secondary) == 'X' else False) if secondary else False
            tertiary_checked = st.checkbox('Tertiary', key=f'{tertiary}_tertiary_checkbox', value='X' if st.session_state.get(tertiary) == 'X' else False) if tertiary else False

            # Set variables based on selections
            if primary_checked:
                st.session_state[primary] = 'X'
            else:
                st.session_state[primary] = ''

            if secondary_checked:
                st.session_state[secondary] = 'X'
            else:
                st.session_state[secondary] = ''

            if tertiary_checked:
                st.session_state[tertiary] = 'X'
            else:
                st.session_state[tertiary] = ''

         # Additional information that may impact learning
        st.session_state.additional_info = st.text_area('Is there any other additional information that may impact on your ability to learn?', value=st.session_state.get('additional_info', ''))

    else:
        st.session_state.has_disability, st.session_state.no_disability = '', 'N'

    # Collect all checkbox variables to check if any are checked
    disability_checked = any([
        st.session_state.get('vision_primary', ''), st.session_state.get('vision_secondary', ''), st.session_state.get('vision_tertiary', ''),
        st.session_state.get('hearing_primary', ''), st.session_state.get('hearing_secondary', ''), st.session_state.get('hearing_tertiary', ''),
        st.session_state.get('mobility_primary', ''), st.session_state.get('mobility_secondary', ''), st.session_state.get('mobility_tertiary', ''),
        st.session_state.get('complex_primary', ''), st.session_state.get('complex_secondary', ''), st.session_state.get('complex_tertiary', ''),
        st.session_state.get('social_primary', ''), st.session_state.get('social_secondary', ''), st.session_state.get('social_tertiary', ''),
        st.session_state.get('mental_primary', ''), st.session_state.get('mental_secondary', ''), st.session_state.get('mental_tertiary', ''),
        st.session_state.get('moderate_primary', ''), st.session_state.get('moderate_secondary', ''), st.session_state.get('moderate_tertiary', ''),
        st.session_state.get('severe_primary', ''), st.session_state.get('severe_secondary', ''), st.session_state.get('severe_tertiary', ''),
        st.session_state.get('dyslexia_primary', ''), st.session_state.get('dyslexia_secondary', ''), st.session_state.get('dyslexia_tertiary', ''),
        st.session_state.get('dyscalculia_primary', ''), st.session_state.get('dyscalculia_secondary', ''), st.session_state.get('dyscalculia_tertiary', ''),
        st.session_state.get('autism_primary', ''), st.session_state.get('autism_secondary', ''), st.session_state.get('autism_tertiary', ''),
        st.session_state.get('aspergers_primary', ''), st.session_state.get('aspergers_secondary', ''), st.session_state.get('aspergers_tertiary', ''),
        st.session_state.get('temporary_primary', ''), st.session_state.get('temporary_secondary', ''), st.session_state.get('temporary_tertiary', ''),
        st.session_state.get('speech_primary', ''), st.session_state.get('speech_secondary', ''), st.session_state.get('speech_tertiary', ''),
        st.session_state.get('physical_primary', ''), st.session_state.get('physical_secondary', ''), st.session_state.get('physical_tertiary', ''),
        st.session_state.get('specific_primary', ''), st.session_state.get('specific_secondary', ''), st.session_state.get('specific_tertiary', ''),
        st.session_state.get('medical_primary', ''), st.session_state.get('medical_secondary', ''), st.session_state.get('medical_tertiary', ''),
        st.session_state.get('other_learning_primary', ''), st.session_state.get('other_learning_secondary', ''), st.session_state.get('other_learning_tertiary', ''),
        st.session_state.get('other_disability_primary', ''), st.session_state.get('other_disability_secondary', ''), st.session_state.get('other_disability_tertiary', ''),
        st.session_state.get('prefer_not_to_say', '')
    ])

    # Other disadvantaged sections
    st.subheader('Other disadvantaged')
    # Initialize ex_offender variable if it doesn't exist
    if 'ex_offender' not in st.session_state:
        st.session_state.ex_offender = 'N'  # Default value

    st.session_state.ex_offender = st.radio('Ex Offender?', ['N', 'Y', 'Choose not to say'], index=['N', 'Y', 'Choose not to say'].index(st.session_state.ex_offender))

    # Initialize ex_offender variables
    st.session_state.ex_offender_y, st.session_state.ex_offender_n, st.session_state.ex_offender_choose_not_to_say = '', '', ''

    # Conditional input for ex_offender option
    if st.session_state.ex_offender == "Y":
        st.session_state.ex_offender_y = 'Y'
    elif st.session_state.ex_offender == "N":
        st.session_state.ex_offender_n = 'N'
    elif st.session_state.ex_offender == "Choose not to say":
        st.session_state.ex_offender_choose_not_to_say = 'Choose not to say'

    # Initialize homeless variable if it doesn't exist
    if 'homeless' not in st.session_state:
        st.session_state.homeless = 'N'  # Default value

    st.session_state.homeless = st.radio('Homeless?', ['N', 'Y', 'Choose not to say'], index=['N', 'Y', 'Choose not to say'].index(st.session_state.homeless))

    # Initialize homeless variables
    st.session_state.homeless_y, st.session_state.homeless_n, st.session_state.homeless_choose_not_to_say = '', '', ''

    # Conditional input for homeless option
    if st.session_state.homeless == "Y":
        st.session_state.homeless_y = 'Y'
    elif st.session_state.homeless == "N":
        st.session_state.homeless_n = 'N'
    elif st.session_state.homeless == "Choose not to say":
        st.session_state.homeless_choose_not_to_say = 'Choose not to say'

    # st.write(disability_checked)

    # Next and Back buttons for navigation
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        # Check if the "disability" is 'Y' and at least one checkbox is checked
        if st.session_state.disability == 'Y' and not disability_checked:
            st.warning("Please select at least one disability type before proceeding.")
        else:
            st.session_state.step = 6
            st.experimental_rerun()

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 4  # Go back to the previous step (Section 3)
        st.experimental_rerun()


elif st.session_state.step == 6:
    st.title("> 5: Referral Source Section")
    # Referral Source Section
    st.header('Referral Source')
    # Creating columns for referral source options
    col1, col2, col3, col4 = st.columns(4)

    # Initialize referral source variables if they do not exist
    if 'internally_sourced' not in st.session_state: st.session_state.internally_sourced = False
    if 'recommendation' not in st.session_state: st.session_state.recommendation = False
    if 'event' not in st.session_state: st.session_state.event = False
    if 'self_referral' not in st.session_state: st.session_state.self_referral = False
    if 'family_friends' not in st.session_state: st.session_state.family_friends = False
    if 'other' not in st.session_state: st.session_state.other = False
    if 'website' not in st.session_state: st.session_state.website = False
    if 'promotional_material' not in st.session_state: st.session_state.promotional_material = False
    if 'jobcentre_plus' not in st.session_state: st.session_state.jobcentre_plus = False
    if 'event_specify' not in st.session_state: st.session_state.event_specify = ''
    if 'other_specify' not in st.session_state: st.session_state.other_specify = ''

    # Adding checkboxes for each referral source option
    with col1:
        st.session_state.internally_sourced = st.checkbox('Internally sourced', value=st.session_state.internally_sourced)
        st.session_state.recommendation = st.checkbox('Recommendation', value=st.session_state.recommendation)
        st.session_state.promotional_material = st.checkbox('Promotional material', value=st.session_state.promotional_material)

    with col2:
        st.session_state.self_referral = st.checkbox('Self Referral', value=st.session_state.self_referral)
        st.session_state.family_friends = st.checkbox('Family/ Friends', value=st.session_state.family_friends)
        st.session_state.event = st.checkbox('Event (please specify)', value=st.session_state.event)

    with col3:
        st.session_state.website = st.checkbox('Website', value=st.session_state.website)
        st.session_state.jobcentre_plus = st.checkbox('JobCentre Plus', value=st.session_state.jobcentre_plus)
        st.session_state.other = st.checkbox('Other (please specify)', value=st.session_state.other)

    # Text inputs for 'Event (please specify)' and 'Other (please specify)' if checked
    if st.session_state.event:
        st.session_state.event_specify = st.text_input('Please specify the event', value=st.session_state.event_specify)

    if st.session_state.other:
        st.session_state.other_specify = st.text_input('Please specify other source', value=st.session_state.other_specify)

    # Additional text input for specifying the referrer
    st.session_state.specify_referral = st.text_input("Please let us know the organization or advisor who referred you to our program, or indicate where you found out about this opportunity. If it was through a job center, please specify its location.", value=st.session_state.specify_referral)

    # Setting 'X' for chosen options
    st.session_state.internally_sourced_val = 'X' if st.session_state.internally_sourced else ''
    st.session_state.recommendation_val = 'X' if st.session_state.recommendation else ''
    st.session_state.event_val = st.session_state.event_specify if st.session_state.event else ''
    st.session_state.self_referral_val = 'X' if st.session_state.self_referral else ''
    st.session_state.family_friends_val = 'X' if st.session_state.family_friends else ''
    st.session_state.other_val = st.session_state.other_specify if st.session_state.other else ''
    st.session_state.website_val = 'X' if st.session_state.website else ''
    st.session_state.promotional_material_val = 'X' if st.session_state.promotional_material else ''
    st.session_state.jobcentre_plus_val = 'X' if st.session_state.jobcentre_plus else ''
    


    # mandatory validation
    # referrall=''
    # if (internally_sourced_val == 'X' or 
    #     recommendation_val == 'X' or 
    #     event_val == 'X' or 
    #     self_referral_val == 'X' or 
    #     family_friends_val == 'X' or 
    #     other_val == 'X' or 
    #     website_val == 'X' or 
    #     promotional_material_val == 'X' or
    #     jobcentre_plus_val == 'X' and
    #     len(specify_referral)>0):
    #     referrall = 'filled'
    #     st.write('LENGTH:', len(specify_referral))
    # mandatory_fields.extend(['p304'])
   
    # Next and Back buttons for navigation
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click with validation
    if next_clicked:
        if (st.session_state.specify_referral):
            st.session_state.step = 7
            st.experimental_rerun()
        else:
            st.warning("Please fill in all fields before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 5  # Go back to the previous step (Section 4)
        st.experimental_rerun()


elif st.session_state.step == 7:
    st.title("> 6: Employment and Monitoring Information Section")

    # Employment and Monitoring Information Section
    st.header('Employment and Monitoring Information')

    # Initialize employment status variables if they do not exist
    if 'unemployed_val' not in st.session_state: st.session_state.unemployed_val = ''
    if 'economically_inactive_val' not in st.session_state: st.session_state.economically_inactive_val = ''
    if 'employed_val' not in st.session_state: st.session_state.employed_val = ''

    # Participant Employment Status
    st.subheader('Participant Employment Status')

    # Set the default value for the radio button based on session state
    st.session_state.employment_status = st.radio(
        "Select your employment status:",
        [
            "Unemployed (looking for work and available to start work) -> go to section A",
            "Economically Inactive (not looking for work and not available to start work) -> Go to section B",
            "Employed (including self-employed) -> go to section C"
        ],
        index=[
            "Unemployed (looking for work and available to start work) -> go to section A",
            "Economically Inactive (not looking for work and not available to start work) -> Go to section B",
            "Employed (including self-employed) -> go to section C"
        ].index(st.session_state.get('employment_status', 
            "Unemployed (looking for work and available to start work) -> go to section A"))  # Default to first option
    )

    # Setting 'X' for chosen employment status
    if st.session_state.employment_status == "Unemployed (looking for work and available to start work) -> go to section A":
        st.session_state.unemployed_val = 'X'
    elif st.session_state.employment_status == "Economically Inactive (not looking for work and not available to start work) -> Go to section B":
        st.session_state.economically_inactive_val = 'X'
    elif st.session_state.employment_status == "Employed (including self-employed) -> go to section C":
        st.session_state.employed_val = 'X'

    st.session_state.up_to_12_months_val, st.session_state.twelve_months_or_longer_val = '-', '-'
    # Section A - Unemployment details
    if "Unemployed" in st.session_state.employment_status:
        st.subheader('Section A - Unemployment details')
        st.text("Where a participant’s employment status is long-term unemployed proof of both unemployment and the length of unemployment must be obtained.")
        
        # Initialize unemployment duration variables if they do not exist
        if 'unemployment_duration' not in st.session_state:
            st.session_state.unemployment_duration = "Up to 12 months"  # Default value

        if 'up_to_12_months_val' not in st.session_state:
            st.session_state.up_to_12_months_val = ''
        if 'twelve_months_or_longer_val' not in st.session_state:
            st.session_state.twelve_months_or_longer_val = ''

        # Unemployment duration selection
        st.session_state.unemployment_duration = st.radio(
            "If you are not working, how long have you been without work?",
            ["Up to 12 months", "12 months or longer"],
            index=["Up to 12 months", "12 months or longer"].index(st.session_state.unemployment_duration)  # Retain state
        )

        # Set 'X' for chosen unemployment duration
        if st.session_state.unemployment_duration == "Up to 12 months":
            st.session_state.up_to_12_months_val = 'X'
            st.session_state.twelve_months_or_longer_val = ''  # Ensure the other value is cleared
        elif st.session_state.unemployment_duration == "12 months or longer":
            st.session_state.twelve_months_or_longer_val = 'X'
            st.session_state.up_to_12_months_val = ''  # Ensure the other value is cleared

        # Evidence of Unemployment Status Section
        st.write("Evidence of unemployment status (for more information look Start-Eligibility Evidence list tab)")

        # Initialize unemployment evidence variable if it does not exist
        if 'unemployment_evidence' not in st.session_state:
            st.session_state.unemployment_evidence = "A Letter or Document from JCP or DWP"  # Default value

        # Unemployment evidence selection
        st.session_state.unemployment_evidence = st.selectbox(
            "Select evidence type:",
            [
                "A Letter or Document from JCP or DWP",
                "A written referral from a careers service",
                "Third Party Verification or Referral form",
                "Other (please specify)"
            ],
            index=[
                "A Letter or Document from JCP or DWP",
                "A written referral from a careers service",
                "Third Party Verification or Referral form",
                "Other (please specify)"
            ].index(st.session_state.unemployment_evidence)  # Retain state
        )

        # Initialize unemployment evidence variables
        st.session_state.jcp_dwp_val = '-' if st.session_state.unemployment_evidence == "A Letter or Document from JCP or DWP" else ''
        st.session_state.careers_service_val = '-' if st.session_state.unemployment_evidence == "A written referral from a careers service" else ''
        st.session_state.third_party_val = '-' if st.session_state.unemployment_evidence == "Third Party Verification or Referral form" else ''
        st.session_state.other_evidence_val = '-' if st.session_state.unemployment_evidence == "Other (please specify)" else ''
        
        # Setting 'X' for chosen evidence type
        if st.session_state.unemployment_evidence == "A Letter or Document from JCP or DWP":
            st.session_state.jcp_dwp_val = 'X'
            uploaded_file = st.file_uploader("Upload Document from JCP or DWP", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file is not None:
                st.session_state.files.append(uploaded_file)
        elif st.session_state.unemployment_evidence == "A written referral from a careers service":
            st.session_state.careers_service_val = 'X'
            uploaded_file = st.file_uploader("Upload written referral from a careers service", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file is not None:
                st.session_state.files.append(uploaded_file)
        elif st.session_state.unemployment_evidence == "Third Party Verification or Referral form":
            st.session_state.third_party_val = 'X'
            uploaded_file = st.file_uploader("Upload Third Party Verification or Referral form", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file is not None:
                st.session_state.files.append(uploaded_file)
        elif st.session_state.unemployment_evidence == "Other (please specify)":
            st.session_state.other_evidence_val = st.text_input("Please specify other evidence", value=st.session_state.other_evidence_val)  # Retain input

    # Section B - Economically Inactive details
    if "Economically Inactive" in st.session_state.employment_status:
        st.subheader('Section B - Economically Inactive details')

        # Economically inactive status selection
        st.session_state.inactive_status = st.radio(
            "The Participant is not employed and does not claim benefits at the time of the enrolment.",
            ["Y", "N"],
            index=["Y", "N"].index(st.session_state.inactive_status) if 'inactive_status' in st.session_state else 0  # Retain state
        )

        # Setting 'X' for chosen inactive status
        st.session_state.inactive_status_val = 'Y' if st.session_state.inactive_status == "Y" else 'N'

        # Evidence input fields
        st.session_state.inactive_evidence_type_val = st.text_input(
            "Type of evidence for Economically Inactive Status including self-declaration statement.",
            value=st.session_state.inactive_evidence_type_val  # Retain value
        )
        
        # Date input for evidence
        if isinstance(st.session_state.get("inactive_evidence_date_val"), str) and st.session_state.get("inactive_evidence_date_val").strip():
            st.session_state.inactive_evidence_date_val = datetime.strptime(st.session_state.get("inactive_evidence_date_val"), "%d-%m-%Y").date()

        st.session_state.inactive_evidence_date_val = st.date_input(
            label="Date of issue of evidence",
            value=st.session_state.get("inactive_evidence_date_val"),
            min_value=date(1900, 1, 1),
            max_value=date.today(),
            help="Choose a date",
            format='DD/MM/YYYY'
        )
        
        if st.session_state.inactive_evidence_date_val != None:
            st.session_state.inactive_evidence_date_val = st.session_state.inactive_evidence_date_val.strftime("%d-%m-%Y")


    # Initialize employment detail variables
    st.session_state.employer_name_val, st.session_state.employer_address_1_val, st.session_state.employer_address_2_val = '', '', ''
    st.session_state.employer_address_3_val, st.session_state.employer_postcode_val, st.session_state.employer_contact_name_val = '', '', ''
    st.session_state.employer_contact_position_val, st.session_state.employer_contact_email_val, st.session_state.employer_contact_phone_val = '', '', ''
    st.session_state.employer_edrs_number_val, st.session_state.living_wage_val, st.session_state.employment_hours_val_0, st.session_state.employment_hours_val_6 = '', '', '', ''
    st.session_state.claiming_benefits_val, st.session_state.sole_claimant_val, st.session_state.benefits_list_val = '', '', ''
    st.session_state.other_benefit_val = ''
    
    # Initialize variables for benefits
    st.session_state.universal_credit_val = ''
    st.session_state.job_seekers_allowance_val = ''
    st.session_state.employment_support_allowance_val = ''
    st.session_state.incapacity_benefit_val = ''
    st.session_state.personal_independence_payment_val = ''

    # Section C - Employment details
    if "Employed" in st.session_state.employment_status:
        st.subheader('Section C - Employment details')

        # Retaining values for text inputs
        st.session_state.employer_name_val = st.text_input("Employer Name", value=st.session_state.get('employer_name_val', ''))
        st.session_state.employer_address_1_val = st.text_input("Employer Address 1", value=st.session_state.get('employer_address_1_val', ''))
        st.session_state.employer_address_2_val = st.text_input("Employer Address 2", value=st.session_state.get('employer_address_2_val', ''))
        st.session_state.employer_address_3_val = st.text_input("Employer Address 3", value=st.session_state.get('employer_address_3_val', ''))
        st.session_state.employer_postcode_val = st.text_input("Employer Postcode", value=st.session_state.get('employer_postcode_val', ''))
        st.session_state.employer_contact_name_val = st.text_input("Main Employer Contact Name", value=st.session_state.get('employer_contact_name_val', ''))
        st.session_state.employer_contact_position_val = st.text_input("Contact Position", value=st.session_state.get('employer_contact_position_val', ''))
        st.session_state.employer_contact_email_val = st.text_input("Contact Email Address", value=st.session_state.get('employer_contact_email_val', ''))
        st.session_state.employer_contact_phone_val = st.text_input("Contact Telephone Number", value=st.session_state.get('employer_contact_phone_val', ''))
        st.session_state.employer_edrs_number_val = st.text_input("Employer EDRS number", value=st.session_state.get('employer_edrs_number_val', ''))

        # Retaining radio button selections
        if 'living_wage' not in st.session_state:
            st.session_state.living_wage = 'N'  # Default value

        st.session_state.living_wage = st.radio(
            "Do you earn more than the National Living Wage of £20,319.00 pa (£10.42ph for 37.5 hrs pw)?",
            ["Y", "N"],
            index=["Y", "N"].index(st.session_state.living_wage)  # Set index based on session state
        )
        st.session_state.living_wage_val = 'Y' if st.session_state.living_wage == "Y" else 'N'

        # Retaining employment hours selections
        if 'employment_hours' not in st.session_state:
            st.session_state.employment_hours = '0-15 hrs per week'  # Default value

        st.session_state.employment_hours = st.radio(
            "Employment Hours (place an X in the applicable box)",
            ["0-15 hrs per week", "16+ hrs per week"],
            index=["0-15 hrs per week", "16+ hrs per week"].index(st.session_state.employment_hours)  # Set index based on session state
        )
        st.session_state.employment_hours_val_0 = 'X' if st.session_state.employment_hours == "0-15 hrs per week" else '-'
        st.session_state.employment_hours_val_6 = 'X' if st.session_state.employment_hours == "16+ hrs per week" else '-'

        st.session_state.job_position = st.text_input("Job Position", value=st.session_state.get('job_position', ''))

        if isinstance(st.session_state.get("job_start_date"), str):
            st.session_state.job_start_date = datetime.strptime(st.session_state.get("job_start_date"), "%d-%m-%Y").date()

        st.session_state.job_start_date = st.date_input(
            label="Job Start Date",  # Label for the field
            value=st.session_state.get("job_start_date"),  # Correctly access job_start_date from session state
            min_value=date(1900, 1, 1),  # Minimum selectable date
            max_value=date.today(),  # Maximum selectable date
            help="Choose a date",  # Tooltip text
            format='DD/MM/YYYY'
        )
        
        if st.session_state.job_start_date != None:
            # Store date in the desired format for further use
            st.session_state.job_start_date = st.session_state.job_start_date.strftime("%d-%m-%Y")

    st.header("Benefits Detail")
    # Radio button for claiming benefits
    if 'claiming_benefits' not in st.session_state:
        st.session_state.claiming_benefits = "N"  # Default value
    st.session_state.claiming_benefits = st.radio("Are you claiming any benefits? If so, please describe below what they are.", ["N", "Y"])

    # Set benefits claiming variable
    st.session_state.claiming_benefits_val = 'Y' if st.session_state.claiming_benefits == "Y" else 'N'

    if st.session_state.claiming_benefits == "Y":
        # Radio button for sole claimant status
        if 'sole_claimant' not in st.session_state:
            st.session_state.sole_claimant = "N"  # Default value
        st.session_state.sole_claimant = st.radio("Are you the sole claimant of the benefit?", ["Y", "N"])

        # Set sole claimant variable
        st.session_state.sole_claimant_val = 'Y' if st.session_state.sole_claimant == "Y" else 'N'

        # Benefits List Section
        if 'benefits_list' not in st.session_state:
            st.session_state.benefits_list = []  # Default to empty list

        # Retain selections in the multiselect
        st.session_state.benefits_list = st.multiselect(
            "Select the benefits you are claiming:",
            [
                "Universal Credit (UC)",
                "Job Seekers Allowance (JSA)",
                "Employment and Support Allowance (ESA)",
                "Incapacity Benefit (or any other sickness related benefit)",
                "Personal Independence Payment (PIP)",
                "Other - please state"
            ],
            default=st.session_state.benefits_list  # Retain selected benefits
        )

        # Initialize benefits variables if not already done
        if 'universal_credit_val' not in st.session_state: st.session_state.universal_credit_val = ''
        if 'job_seekers_allowance_val' not in st.session_state: st.session_state.job_seekers_allowance_val = ''
        if 'employment_support_allowance_val' not in st.session_state: st.session_state.employment_support_allowance_val = ''
        if 'incapacity_benefit_val' not in st.session_state: st.session_state.incapacity_benefit_val = ''
        if 'personal_independence_payment_val' not in st.session_state: st.session_state.personal_independence_payment_val = ''
        if 'other_benefit_val' not in st.session_state: st.session_state.other_benefit_val = ''

        # Update the respective variables based on user selections
        st.session_state.universal_credit_val = 'X' if "Universal Credit (UC)" in st.session_state.benefits_list else ''
        st.session_state.job_seekers_allowance_val = 'X' if "Job Seekers Allowance (JSA)" in st.session_state.benefits_list else ''
        st.session_state.employment_support_allowance_val = 'X' if "Employment and Support Allowance (ESA)" in st.session_state.benefits_list else ''
        st.session_state.incapacity_benefit_val = 'X' if "Incapacity Benefit (or any other sickness related benefit)" in st.session_state.benefits_list else ''
        st.session_state.personal_independence_payment_val = 'X' if "Personal Independence Payment (PIP)" in st.session_state.benefits_list else ''

        # Handle "Other - please state" input
        if "Other - please state" in st.session_state.benefits_list:
            st.session_state.other_benefit_val = st.text_input("Please state other benefit", value=st.session_state.get('other_benefit_val', ''))
        else:
            st.session_state.other_benefit_val = ''  # Reset if not selected

        # Check if benefit_claim_date_val is a string and convert it back to a date object
        if isinstance(st.session_state.get("benefit_claim_date_val"), str):
            st.session_state.benefit_claim_date_val = datetime.strptime(st.session_state.get("benefit_claim_date_val"), "%d-%m-%Y").date()

        # Date of Benefit Claim Date
        st.session_state.benefit_claim_date_val = st.date_input(
            label="From what date has the above claim been in effect?",  # Label for the field
            value=st.session_state.get("benefit_claim_date_val", datetime.now().date()),  # Correctly access from session state
            min_value=date(1900, 1, 1),  # Minimum selectable date
            max_value=date.today(),  # Maximum selectable date
            help="Choose a date",  # Tooltip text
            format='DD/MM/YYYY'
        )
        if st.session_state.benefit_claim_date_val != None:
            # Store date in the desired format for further use
            st.session_state.benefit_claim_date_val = st.session_state.benefit_claim_date_val.strftime("%d-%m-%Y")

        # Check if benefit_claim_date_val is not selected
        if st.session_state.benefit_claim_date_val:
            st.session_state.benefit_claim_date_val = st.session_state.benefit_claim_date_val.strftime("%d-%m-%Y")
        else:
            st.warning("Please choose Benefit Claim Date.")
            st.stop()




    # # Detailed Learning Plan Section
    # st.header('Detailed Learning Plan')

    # qualification_reference = st.text_input("Qualification Reference")
    # region_of_work = st.text_input("Region of Work")
    # qualification_course_title = st.text_input("Qualification/Course/Unit Title/Non-Regulated activity")
    # awarding_body = st.text_input("Awarding Body")

    # GLH = st.text_input("GLH")

    # benefit_to_you = st.text_area("What is the benefit to you in completing this learning aim? Please be specific")

    # planned_start_date = st.date_input("Planned Start Date")
    # planned_end_date = st.date_input("Planned End Date", help="Note: Actual End Date to be recorded on 'Outcome and Progression' form at the end of the programme")
    # delivery_postcode = st.text_input("Delivery Postcode")
    # date_of_first_review = st.date_input("Date of first review")

    # st.subheader("Progression - Indicate below the progression planned for this participant when they have completed all training")
    # progression_options = st.multiselect(
    #     "Select progression options:",
    #     [
    #         "Progression within Work",
    #         "Progression into Further Education or Training",
    #         "Progression to Apprenticeship",
    #         "Progression into employment"
    #     ]
    # )

    # progression_aim = st.text_area("Please detail your progression aim")

    # st.subheader("Social Outcomes - How do you rate yourself now out of 5 for the below. 5= Great 1= Poor")

    # health_and_well_being = st.slider("Health and well being", 1, 5, 1)
    # social_integration = st.slider("Social integration", 1, 5, 1)
    # learner_self_efficacy = st.slider("Learner self-efficacy", 1, 5, 1)
    # participation_in_volunteering = st.slider("Participation in volunteering", 1, 5, 1)








    # st.header('Eligibility Check')

    # st.text("""
    #     Evidence CANNOT be accepted that has been entered at a later date than Actual End Date of the start aim.
    #     Evidence must be present for ALL 4 (EO1,2,3,4) of the below eligibility checks. Original documentation must have been witnessed by the Provider and preferably copies made as evidence in case of future audits.
    #     For list of ALL acceptable supporting documents check 'Start-Eligibility Evidence list'
    #     """)

    # st.text("""
    #     UK, EEA Nationals and Non-EEA Nationals

    #     EEA Countries (as of 27/01/2021): 
    #     Austria, Belgium, Bulgaria, Croatia, Republic of Cyprus, Czech Republic, Denmark, Estonia, Finland, France, Germany, Greece, Hungary, Ireland, Italy, Latvia, Lithuania, Luxembourg, Malta, Netherlands, Poland, Portugal, Romania, Slovakia, Slovenia, Spain, Sweden, Iceland, Liechtenstein, Norway.

    #     Switzerland is not an EU or EEA member but is part of the single market. This means Swiss nationals have the same rights to live and work in the UK as other EEA nationals.

    #     “Irish citizens in the UK hold a unique status under each country’s national law. You do not need permission to enter or remain in the UK, including a visa, any form of residence permit or employment permit”. Quote taken from below link:
    #     https://www.gov.uk/government/publications/common-travel-area-guidance/common-travel-area-guidance

    #     Non-EEA nationals who hold leave to enter or leave to remain with a permission to work (including status under the EUSS where they are an eligible family member of an EEA national) are also eligible for ESF support whilst in the UK.
    #     """)


    st.header('E01: Right to Live and Work in the UK')

    # Initialize session state variables if they do not exist
    if 'resident' not in st.session_state:
        st.session_state.resident = 'Yes'  # Default value
    if 'resident_y' not in st.session_state:
        st.session_state.resident_y = ''
    if 'resident_n' not in st.session_state:
        st.session_state.resident_n = ''
    if 'country_of_birth' not in st.session_state:
        st.session_state.country_of_birth = ''
    if 'years_in_uk' not in st.session_state:
        st.session_state.years_in_uk = 0  # Default to 0
    if 'british_or_not' not in st.session_state:
        st.session_state.british_or_not = 'Yes'  # Default value

    # Resident radio selection
    st.session_state.resident = st.radio(
        'Have you been resident in the UK/EEA for the previous 3 years?',
        ('Yes', 'No'),
        index=0 if st.session_state.resident == 'Yes' else 1  # Set the default based on session state
    )

    if st.session_state.resident == 'Yes':
        st.session_state.resident_y = 'X'
        st.session_state.resident_n = ''
    else:
        st.session_state.resident_n = 'X'
        st.session_state.resident_y = ''

    # Input fields for country of birth and years in the UK
    st.session_state.country_of_birth = st.text_input('Country of Birth:', value=st.session_state.country_of_birth)
    st.session_state.years_in_uk = st.number_input('How many years have you lived in the UK?', min_value=0, value=st.session_state.years_in_uk)

    # Variable initialization
    if 'hold_settled_status' not in st.session_state: st.session_state.hold_settled_status = ''
    if 'hold_pre_settled_status' not in st.session_state: st.session_state.hold_pre_settled_status = ''
    if 'hold_leave_to_remain' not in st.session_state: st.session_state.hold_leave_to_remain = ''
    if 'not_nationality' not in st.session_state: st.session_state.not_nationality = ''
    if 'passport_non_eu' not in st.session_state: st.session_state.passport_non_eu = ''
    if 'letter_uk_immigration' not in st.session_state: st.session_state.letter_uk_immigration = ''
    if 'passport_endorsed' not in st.session_state: st.session_state.passport_endorsed = ''
    if 'identity_card' not in st.session_state: st.session_state.identity_card = ''
    if 'country_of_issue' not in st.session_state: st.session_state.country_of_issue = ''
    if 'id_document_reference_number' not in st.session_state: st.session_state.id_document_reference_number = ''
    if 'e01_date_of_issue' not in st.session_state: st.session_state.e01_date_of_issue = ''
    if 'e01_date_of_expiry' not in st.session_state: st.session_state.e01_date_of_expiry = ''
    if 'e01_additional_notes' not in st.session_state: st.session_state.e01_additional_notes = ''

    # Create a radio button for the Yes/No question
    st.session_state.british_or_not = st.radio(
        'Are you a UK OR Irish National OR European Economic Area (EEA) National?',
        ('Yes', 'No'),
        index=0 if st.session_state.british_or_not == 'Yes' else 1  # Set the default based on session state
    )


    # Initialize nationality and document variables
    if 'nationality' not in st.session_state: 
        st.session_state.nationality = ''
    if 'full_uk_passport' not in st.session_state: 
        st.session_state.full_uk_passport = ''
    if 'full_eu_passport' not in st.session_state: 
        st.session_state.full_eu_passport = ''
    if 'national_identity_card' not in st.session_state: 
        st.session_state.national_identity_card = ''
    if 'selected_option_nationality' not in st.session_state: 
        st.session_state.selected_option_nationality = 'Full UK Passport'  # Default value

    if st.session_state.british_or_not == 'Yes':
        st.session_state.nationality = st.text_input('Nationality', value=st.session_state.nationality)
        
        options = [
            'Full UK Passport',
            'Full EU Member Passport (must be in date - usually 10 years)',
            'National Identity Card (EU)'
        ]
        
        st.session_state.selected_option_nationality = st.radio("Select the type of document:", options, index=options.index(st.session_state.selected_option_nationality))

        if st.session_state.selected_option_nationality == options[0]:
            st.session_state.full_uk_passport, st.session_state.full_eu_passport, st.session_state.national_identity_card = 'X', '', ''
            st.text('Please upload a copy of your Full UK Passport')
            uploaded_file = st.file_uploader("Upload Full UK Passport", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file is not None:
                st.session_state.files.append(uploaded_file)
            uploaded_file_2 = st.file_uploader("Optional - Upload Back Side of Document", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file_2 is not None:
                st.session_state.files.append(uploaded_file_2)

        elif st.session_state.selected_option_nationality == options[1]:
            st.session_state.full_uk_passport, st.session_state.full_eu_passport, st.session_state.national_identity_card = '', 'X', ''
            st.text('Please upload a copy of your Full EU Member Passport')
            uploaded_file = st.file_uploader("Upload Full EU Member Passport", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file is not None:
                st.session_state.files.append(uploaded_file)
            uploaded_file_2 = st.file_uploader("Optional - Upload Back Side of Document", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file_2 is not None:
                st.session_state.files.append(uploaded_file_2)

        elif st.session_state.selected_option_nationality == options[2]:
            st.session_state.full_uk_passport, st.session_state.full_eu_passport, st.session_state.national_identity_card = '', '', 'X'
            st.text('Please upload a copy of your National Identity Card (EU)')
            uploaded_file = st.file_uploader("Upload National Identity Card (EU)", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file is not None:
                st.session_state.files.append(uploaded_file)
            uploaded_file_2 = st.file_uploader("Optional - Upload Back Side of Document", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file_2 is not None:
                st.session_state.files.append(uploaded_file_2)

        if st.session_state.selected_option_nationality in [options[1], options[2]]:
            st.text(
                'In order to be eligible for ESF funding, EEA Nationals must meet one of the following conditions'
            )
            conditions = [
                'a. Hold settled status granted under the EU Settlement Scheme (EUSS)',
                'b. Hold pre-settled status granted under the European Union Settlement Scheme (EUSS)',
                'c. Hold leave to remain with permission to work granted under the new Points Based Immigration System'
            ]

            # Initially set the radio button without any selection
            if 'settled_status' not in st.session_state:
                st.session_state.settled_status = None  # Default to no selection

            st.session_state.settled_status = st.radio("Select your status:", options=conditions, index=None if st.session_state.settled_status is None else conditions.index(st.session_state.settled_status))

            # Check if no selection is made
            if not st.session_state.settled_status:
                st.warning("Please select your status before proceeding.")
                st.stop()

            if st.session_state.settled_status == conditions[0]:
                st.session_state.hold_settled_status, st.session_state.hold_pre_settled_status, st.session_state.hold_leave_to_remain = 'X', '', ''
                st.text('Please upload your share code which is accessible from the following link:')
                uploaded_file = st.file_uploader("https://www.gov.uk/check-immigration-status", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file is not None:
                    st.session_state.files.append(uploaded_file)
                uploaded_file_3 = st.file_uploader("Optional - Upload Back Side of Document ", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file_3 is not None:
                    st.session_state.files.append(uploaded_file_3)

            elif st.session_state.settled_status == conditions[1]:
                st.session_state.hold_settled_status, st.session_state.hold_pre_settled_status, st.session_state.hold_leave_to_remain = '', 'X', ''
                st.text('Please upload your share code which is accessible from the following link:')
                uploaded_file = st.file_uploader("https://www.gov.uk/check-immigration-status", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file is not None:
                    st.session_state.files.append(uploaded_file)
                uploaded_file_3 = st.file_uploader("Optional - Upload Back Side of Document  ", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file_3 is not None:
                    st.session_state.files.append(uploaded_file_3)

            elif st.session_state.settled_status == conditions[2]:
                st.session_state.hold_settled_status, st.session_state.hold_pre_settled_status, st.session_state.hold_leave_to_remain = '', '', 'X'
                st.text('Please upload your share code which is accessible from the following link:')
                uploaded_file = st.file_uploader("https://www.gov.uk/check-immigration-status", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file is not None:
                    st.session_state.files.append(uploaded_file)
                uploaded_file_3 = st.file_uploader("Optional - Upload Back Side of Document   ", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file_3 is not None:
                    st.session_state.files.append(uploaded_file_3)

        else:
            st.session_state.not_nationality = st.text_input('Nationality ', value=st.session_state.not_nationality)
            st.session_state.passport_non_eu_checked = st.checkbox(
                'Passport from non-EU member state (must be in date) AND any of the below a, b, or c',
                value='X' if st.session_state.passport_non_eu_checked == 'X' else False  # Retain checkbox state
            )
            if st.session_state.passport_non_eu_checked:
                st.session_state.passport_non_eu = 'X'
                st.text('Please upload a copy of your non-EU Passport')
                uploaded_file = st.file_uploader("Upload Non-EU Passport", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file is not None:
                    st.session_state.files.append(uploaded_file)
                uploaded_file_2 = st.file_uploader("Optional - Upload Back Side of Document", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file_2 is not None:
                    st.session_state.files.append(uploaded_file_2)
            else:
                st.session_state.passport_non_eu = ''

            document_options = [
                "a. Letter from the UK Immigration and Nationality Directorate granting indefinite leave to remain (settled status)",
                "b. Passport either endorsed 'indefinite leave to remain' – (settled status) or includes work or residency permits or visa stamps (unexpired) and all related conditions met; add details below",
                "c. Some non-EEA nationals have an Identity Card (Biometric Permit) issued by the Home Office in place of a visa, confirming the participant’s right to stay, work or study in the UK – these cards are acceptable"
            ]

            # Initially set the radio button without any selection
            if 'document_type' not in st.session_state:
                st.session_state.document_type = None  # Default to no selection

            st.session_state.document_type = st.radio("Select the type of document:", options=document_options, index=None if st.session_state.document_type is None else document_options.index(st.session_state.document_type))

            # Check if no selection is made
            if not st.session_state.document_type:
                st.warning("Please select the type of document before proceeding.")
                st.stop()
            st.session_state.letter_uk_immigration, st.session_state.passport_endorsed, st.session_state.identity_card = '', '', ''

            if st.session_state.document_type == document_options[0]:
                st.session_state.letter_uk_immigration, st.session_state.passport_endorsed, st.session_state.identity_card = 'X', '', ''
                st.text('Please upload your Letter from the UK Immigration and Nationality Directorate')
                uploaded_file = st.file_uploader("Upload Letter from UK Immigration and Nationality Directorate", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file is not None:
                    st.session_state.files.append(uploaded_file)
                uploaded_file_4 = st.file_uploader("Optional - Upload Back Side of Document ", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file_4 is not None:
                    st.session_state.files.append(uploaded_file_4)

            elif st.session_state.document_type == document_options[1]:
                st.session_state.letter_uk_immigration, st.session_state.passport_endorsed, st.session_state.identity_card = '', 'X', ''
                st.text('Please upload your endorsed passport')
                uploaded_file = st.file_uploader("Upload Endorsed Passport", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file is not None:
                    st.session_state.files.append(uploaded_file)
                uploaded_file_4 = st.file_uploader("Optional - Upload Back Side of Document  ", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file_4 is not None:
                    st.session_state.files.append(uploaded_file_4)

            elif st.session_state.document_type == document_options[2]:
                st.session_state.letter_uk_immigration, st.session_state.passport_endorsed, st.session_state.identity_card = '', '', 'X'
                st.text('Please upload your Identity Card (Biometric Permit)')
                uploaded_file = st.file_uploader("Upload Identity Card (Biometric Permit)", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file is not None:
                    st.session_state.files.append(uploaded_file)
                uploaded_file_4 = st.file_uploader("Optional - Upload Back Side of Document   ", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file_4 is not None:
                    st.session_state.files.append(uploaded_file_4)

    # Initialize session state variables for the new fields if they do not exist
    if 'country_of_issue' not in st.session_state: 
        st.session_state.country_of_issue = ''
    if 'id_document_reference_number' not in st.session_state: 
        st.session_state.id_document_reference_number = ''
    if 'e01_date_of_issue' not in st.session_state: 
        st.session_state.e01_date_of_issue = datetime(2000, 1, 1)
    if 'e01_date_of_expiry' not in st.session_state: 
        st.session_state.e01_date_of_expiry = datetime(2000, 1, 1)
    if 'e01_additional_notes' not in st.session_state: 
        st.session_state.e01_additional_notes = ''

    # Input fields for the country of issue and ID document reference number
    st.session_state.country_of_issue = st.text_input('Country of issue', value=st.session_state.country_of_issue)
    st.session_state.id_document_reference_number = st.text_input('ID Document Reference Number', value=st.session_state.id_document_reference_number)

    # Check if the date is a non-empty string, then parse it
    if isinstance(st.session_state.get("e01_date_of_issue"), str) and st.session_state.get("e01_date_of_issue").strip():
        st.session_state.e01_date_of_issue = datetime.strptime(st.session_state.get("e01_date_of_issue"), "%d-%m-%Y").date()

    # Date of Issue input with optional handling
    st.session_state.e01_date_of_issue = st.date_input(
        label="Date of Issue",
        value=st.session_state.e01_date_of_issue or None,
        min_value=date(1900, 1, 1),
        max_value=date(2025, 12, 31),
        help="Choose a date",
        format='DD/MM/YYYY'
    )

    # If a date is selected, store it as a formatted string
    if st.session_state.e01_date_of_issue:
        st.session_state.e01_date_of_issue = st.session_state.e01_date_of_issue.strftime("%d-%m-%Y")

    if isinstance(st.session_state.get("e01_date_of_expiry"), str) and st.session_state.get("e01_date_of_expiry").strip():
        st.session_state.e01_date_of_expiry = datetime.strptime(
            st.session_state.get("e01_date_of_expiry"), "%d-%m-%Y"
        ).date()
    else:
        # Leave it as None if there's no valid date
        st.session_state.e01_date_of_expiry = None

    # Date of Expiry
    st.session_state.e01_date_of_expiry = st.date_input(
        label="Date of Expiry",
        value=st.session_state.e01_date_of_expiry,  # Retrieve the stored date
        min_value=date(1900, 1, 1),  # Minimum selectable date
        max_value=date(2050, 12, 31),  # Maximum selectable date
        help="Choose a date",  # Tooltip text
        format='DD/MM/YYYY'
    )
    if st.session_state.e01_date_of_expiry != None:
        st.session_state.e01_date_of_expiry = st.session_state.e01_date_of_expiry.strftime("%d-%m-%Y")

    # Additional Notes
    st.session_state.e01_additional_notes = st.text_area('Use this space for additional notes where relevant (type of Visa, restrictions, expiry etc.)', value=st.session_state.e01_additional_notes)
    
    
    # Ensure all relevant session state variables are initialized
    if 'full_passport_eu' not in st.session_state: 
        st.session_state.full_passport_eu = False
    if 'national_id_card_eu' not in st.session_state: 
        st.session_state.national_id_card_eu = False
    if 'firearms_certificate' not in st.session_state: 
        st.session_state.firearms_certificate = False
    if 'birth_adoption_certificate' not in st.session_state: 
        st.session_state.birth_adoption_certificate = False
    if 'e02_drivers_license' not in st.session_state: 
        st.session_state.e02_drivers_license = False
    if 'edu_institution_letter' not in st.session_state: 
        st.session_state.edu_institution_letter = False
    if 'e02_employment_contract' not in st.session_state: 
        st.session_state.e02_employment_contract = False
    if 'state_benefits_letter' not in st.session_state: 
        st.session_state.state_benefits_letter = False
    if 'pension_statement' not in st.session_state: 
        st.session_state.pension_statement = False
    if 'northern_ireland_voters_card' not in st.session_state: 
        st.session_state.northern_ireland_voters_card = False
    if 'e02_other_evidence_text' not in st.session_state: 
        st.session_state.e02_other_evidence_text = ''

    st.header('E02: Proof of Age')

    # Adding checkboxes with upload functionality
    st.session_state.full_passport_eu = add_checkbox_with_upload('Full Passport (EU Member State)', 'full_passport_eu')
    st.session_state.national_id_card_eu = add_checkbox_with_upload('National ID Card (EU)', 'national_id_card_eu')
    st.session_state.firearms_certificate = add_checkbox_with_upload('Firearms Certificate/Shotgun Licence', 'firearms_certificate')
    st.session_state.birth_adoption_certificate = add_checkbox_with_upload('Birth/Adoption Certificate', 'birth_adoption_certificate')
    st.session_state.e02_drivers_license = add_checkbox_with_upload('Drivers Licence (photo card)', 'e02_drivers_license')
    st.session_state.edu_institution_letter = add_checkbox_with_upload('Letter from Educational Institution* (showing DOB)', 'edu_institution_letter')
    st.session_state.e02_employment_contract = add_checkbox_with_upload('Employment Contract/Pay Slip (showing DOB)', 'e02_employment_contract')
    st.session_state.state_benefits_letter = add_checkbox_with_upload('State Benefits Letter* (showing DOB)', 'state_benefits_letter')
    st.session_state.pension_statement = add_checkbox_with_upload('Pension Statement* (showing DOB)', 'pension_statement')
    st.session_state.northern_ireland_voters_card = add_checkbox_with_upload('Northern Ireland voters card', 'northern_ireland_voters_card')

    # Other evidence input
    st.session_state.e02_other_evidence_text = st.text_input('Other Evidence: Please state type', value=st.session_state.e02_other_evidence_text)

    # Validation for the last 3 months
    st.session_state.current_date = date.today()
    st.session_state.three_months_ago = st.session_state.current_date - timedelta(days=90)

    if isinstance(st.session_state.get("e02_date_of_issue"), str) and st.session_state.get("e02_date_of_issue").strip():
        st.session_state.e02_date_of_issue = datetime.strptime(st.session_state.get("e02_date_of_issue"), "%d-%m-%Y").date()

    # Date of Issue input
    st.session_state.e02_date_of_issue = st.date_input(
        label="Date of Issue of evidence",
        value=st.session_state.e02_date_of_issue,  # Use the stored date
        min_value=date(1900, 1, 1),  # Minimum selectable date
        max_value=date.today(),  # Maximum selectable date
        help="Choose a date",  # Tooltip text
        format='DD/MM/YYYY'
    )

    if st.session_state.e02_date_of_issue != None:
        st.session_state.e02_date_of_issue = st.session_state.e02_date_of_issue.strftime("%d-%m-%Y")

    # # Check if the selected date is within the last three months
    # if e02_date_of_issue < three_months_ago:
    #     st.warning("The date of issue is not within the last 3 months. Please select a valid date.")
    #     st.stop()
    # st.success("The date of issue is within the last 3 months.")


    # Ensure the date is in the correct format and retained properly
    if isinstance(st.session_state.e02_date_of_issue, str):
        st.session_state.e02_date_of_issue = datetime.strptime(st.session_state.e02_date_of_issue, "%d-%m-%Y").date()

    # Validation for mandatory field
    documents = [
        st.session_state.full_passport_eu,
        st.session_state.national_id_card_eu,
        st.session_state.firearms_certificate,
        st.session_state.birth_adoption_certificate,
        st.session_state.e02_drivers_license,
        st.session_state.edu_institution_letter,
        st.session_state.e02_employment_contract,
        st.session_state.state_benefits_letter,
        st.session_state.pension_statement,
        st.session_state.northern_ireland_voters_card,
    ]

    # Check if at least one of the variables is 'X' or if e02_other_evidence_text is not empty
    if any(doc == 'X' for doc in documents) or st.session_state.e02_other_evidence_text != '':
        st.session_state.e02_filled = 'Filled'
    else:
        st.session_state.e02_filled = ''

    # Update the mandatory fields list if needed
    # mandatory_fields.extend(['p301'])

    st.header('E03: Proof of Residence (must show the address recorded on ILP) *within the last 3 months')

    # Initialize checkbox variables
    st.session_state.e03_drivers_license = add_checkbox_with_upload('Drivers Licence (photo card)', 'e03_drivers_license')
    st.session_state.bank_statement = add_checkbox_with_upload('Bank Statement *', 'bank_statement')
    st.session_state.e03_pension_statement = add_checkbox_with_upload('Pension Statement*', 'e03_pension_statement')
    st.session_state.mortgage_statement = add_checkbox_with_upload('Mortgage Statement*', 'mortgage_statement')
    st.session_state.utility_bill = add_checkbox_with_upload('Utility Bill* (excluding mobile phone)', 'utility_bill')
    st.session_state.council_tax_statement = add_checkbox_with_upload('Council Tax annual statement or monthly bill*', 'council_tax_statement')
    st.session_state.electoral_role_evidence = add_checkbox_with_upload('Electoral Role registration evidence*', 'electoral_role_evidence')
    st.session_state.homeowner_letter = add_checkbox_with_upload('Letter/confirmation from homeowner (family/lodging)', 'homeowner_letter')

    # Initialize text input for other evidence
    st.session_state.e03_other_evidence_text = st.text_input('Other Evidence: Please state type ', value=st.session_state.get('e03_other_evidence_text', ''))

    # Convert `e03_date_of_issue` to a date object if it's a string
    if isinstance(st.session_state.get("e03_date_of_issue"), str):
        st.session_state.e03_date_of_issue = datetime.strptime(st.session_state.get("e03_date_of_issue"), "%d-%m-%Y").date()

    # Date input for evidence
    st.session_state.e03_date_of_issue = st.date_input(
        label="Date of Issue evidence",
        value=st.session_state.get("e03_date_of_issue"),  # Default date
        min_value=date(1900, 1, 1),  # Minimum selectable date
        max_value=date.today(),  # Maximum selectable date
        help="Choose a date",  # Tooltip text
        format='DD/MM/YYYY'
    )

    if st.session_state.e03_date_of_issue != None:
        st.session_state.e03_date_of_issue = st.session_state.e03_date_of_issue.strftime("%d-%m-%Y")

    # Check if the selected date is within the last three months

# Check if the selected date is within the last three months
    if st.session_state.e03_date_of_issue is not None and st.session_state.e03_date_of_issue < st.session_state.three_months_ago:
        st.warning("The date of issue is not within the last 3 months. Please select a valid date.")
        st.stop()
        st.success("The date of issue is within the last 3 months.")
        st.session_state.e03_date_of_issue = st.session_state.e03_date_of_issue.strftime("%d-%m-%Y")

    # Validation for mandatory fields
    documents = [
        st.session_state.e03_drivers_license,
        st.session_state.bank_statement,
        st.session_state.e03_pension_statement,
        st.session_state.mortgage_statement,
        st.session_state.utility_bill,
        st.session_state.council_tax_statement,
        st.session_state.electoral_role_evidence,
        st.session_state.homeowner_letter,
    ]

    # Check if at least one of the variables is 'X' or if e03_other_evidence_text is not empty
    if any(doc == 'X' for doc in documents) or st.session_state.e03_other_evidence_text != '':
        st.session_state.e03_filled = 'Filled'
    else:
        st.session_state.e03_filled = ''

    # Update the mandatory fields list if needed
    # mandatory_fields.extend(['p302'])

    st.header('E04: Employment Status (please select one option from below and take a copy)')

    # Initialize employment status variables
    st.session_state.latest_payslip = st.session_state.get('latest_payslip', '-')
    st.session_state.e04_employment_contract = st.session_state.get('e04_employment_contract', '-')
    st.session_state.confirmation_from_employer = st.session_state.get('confirmation_from_employer', '-')
    st.session_state.redundancy_notice = st.session_state.get('redundancy_notice', '-')
    st.session_state.sa302_declaration = st.session_state.get('sa302_declaration', '-')
    st.session_state.ni_contributions = st.session_state.get('ni_contributions', '-')
    st.session_state.business_records = st.session_state.get('business_records', '-')
    st.session_state.companies_house_records = st.session_state.get('companies_house_records', '-')
    st.session_state.other_evidence_employed = st.session_state.get('other_evidence_employed', '-')
    st.session_state.unemployed = st.session_state.get('unemployed', '-')

    main_options = [
        'a. Latest Payslip (maximum 3 months prior to start date)',
        'b. Employment Contract',
        'c. Confirmation from the employer that the Participant is currently employed by them which must detail: Participant full name, contracted hours, start date AND date of birth or NINO',
        'd. Redundancy consultation or notice (general notice to group of staff or individual notifications) At risk of Redundancy only',
        'e. Self-employed',
        'f. Other evidence as listed in the \'Start-Eligibility Evidence list\' under Employed section - State below',
        'g. Unemployed (complete the Employment section in ILP form)'
    ]

    # Set selected option from the previous selection
    st.session_state.selected_main_option = st.radio("Select an employment status or document:", main_options, index=main_options.index(st.session_state.get('selected_main_option', main_options[0])))

    if st.session_state.selected_main_option == main_options[0]:
        st.session_state.latest_payslip = 'X'
        handle_file_upload('Latest Payslip (maximum 3 months prior to start date)')

        # Validation for the date of issue
        st.session_state.current_date = date.today()
        st.session_state.three_months_ago = st.session_state.current_date - timedelta(days=90)

        if isinstance(st.session_state.get("e04_date_of_issue"), str):
            st.session_state.e04_date_of_issue = datetime.strptime(st.session_state.get("e04_date_of_issue"), "%d-%m-%Y").date()

        st.session_state.e04_date_of_issue = st.date_input(
            label="Date of Issue of evidence ",
            value=st.session_state.get('e04_date_of_issue', date.today()),  # Default date
            min_value=date(1900, 1, 1),  # Minimum selectable date
            max_value=date.today(),  # Maximum selectable date
            help="Choose a date",  # Tooltip text
            format='DD/MM/YYYY'
        )

        if st.session_state.e04_date_of_issue != None:
            st.session_state.e04_date_of_issue = st.session_state.e04_date_of_issue.strftime("%d-%m-%Y")

        if st.session_state.e04_date_of_issue < st.session_state.three_months_ago:
            st.warning("The date of issue is not within the last 3 months. Please select a valid date.")
            st.stop()
        st.success("The date of issue is within the last 3 months.")
        st.session_state.e04_date_of_issue = st.session_state.e04_date_of_issue.strftime("%d-%m-%Y")

    elif st.session_state.selected_main_option == main_options[1]:
        st.session_state.e04_employment_contract = 'X'
        handle_file_upload('Employment Contract')
    elif st.session_state.selected_main_option == main_options[2]:
        st.session_state.confirmation_from_employer = 'X'
        handle_file_upload('Confirmation from the employer')
    elif st.session_state.selected_main_option == main_options[3]:
        st.session_state.redundancy_notice = 'X'
        handle_file_upload('Redundancy consultation or notice')
    elif st.session_state.selected_main_option == main_options[4]:
        self_employed_options = [
            "HMRC 'SA302' self-assessment tax declaration, with acknowledgement of receipt (within last 12 months)",
            'Records to show actual payment of Class 2 National Insurance Contributions (within last 12 months)',
            'Business records in the name of the business - evidence that a business has been established and is active / operating (within last 12 months)',
            'If registered as a Limited company: Companies House records / listed as Company Director (within last 12 months)'
        ]
        st.session_state.selected_self_employed_option = st.radio("Select self-employed evidence:", self_employed_options, index=self_employed_options.index(st.session_state.get('selected_self_employed_option', self_employed_options[0])))

        if st.session_state.selected_self_employed_option == self_employed_options[0]:
            st.session_state.sa302_declaration = 'X'
            handle_file_upload("HMRC 'SA302' self-assessment tax declaration")
        elif st.session_state.selected_self_employed_option == self_employed_options[1]:
            st.session_state.ni_contributions = 'X'
            handle_file_upload('Records of Class 2 National Insurance Contributions')
        elif st.session_state.selected_self_employed_option == self_employed_options[2]:
            st.session_state.business_records = 'X'
            handle_file_upload('Business records')
        elif st.session_state.selected_self_employed_option == self_employed_options[3]:
            st.session_state.companies_house_records = 'X'
            handle_file_upload('Companies House records')
    elif st.session_state.selected_main_option == main_options[5]:
        st.session_state.other_evidence_employed = 'X'
        handle_file_upload("Other evidence as listed in the 'Start-Eligibility Evidence list'")
    elif st.session_state.selected_main_option == main_options[6]:
        st.session_state.unemployed = 'X'
        handle_file_upload('Unemployed (complete the Employment section in ILP form)')

    # Navigation buttons
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        
        st.session_state.job_start_date = st.session_state.job_start_date.strftime("%d-%m-%Y")


        # if (st.session_state.country_of_issue and st.session_state.id_document_reference_number and st.session_state.e01_additional_notes):
        st.session_state.step = 8
        st.experimental_rerun()
        # else:
            # st.warning("Please fill 'Country of issue' and 'ID Document Reference Number' and 'Additional Note'")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 6  # Go back to the previous step (Section 5)
        st.experimental_rerun()

elif st.session_state.step == 8:
    st.title("> 7: Details of Qualification or Training")

    st.header('Details of Qualification or Training')
  
    st.session_state.qualification_or_training = st.radio(
    'Are you currently undertaking a qualification or training?',
    ['No', 'Yes'])

    if st.session_state.qualification_or_training=='Yes':
        st.session_state.qualification_or_training_y, st.session_state.qualification_or_training_n = 'Y', ''

        st.session_state.course_details = st.text_area('Course Details',
                                      'Enter details of the course')
        st.session_state.funding_details = st.text_area(
            'Funding Details', 'Enter details of how the course is funded')
    else:
        st.session_state.qualification_or_training_y, st.session_state.qualification_or_training_n = '', 'N'
        st.session_state.course_details, st.session_state.funding_details = '', ''
        st.write(
            'You answered "No" to currently undertaking a qualification or training.'
        )

    st.header('Evidenced Qualification Levels')


    st.subheader('Participant self declaration of highest qualification level')
    participant_options = [
        'Below Level 1', 'Level 1', 'Level 2', 'Full Level 2', 'Level 3', 'Full Level 3', 'Level 4',
        'Level 5', 'Level 6', 'Level 7 and above', 'No Qualifications'
    ]


    st.session_state.participant_declaration = st.radio('', participant_options)


    st.session_state.p58 = '-'
    st.session_state.p59 = '-'
    st.session_state.p60 = '-'
    st.session_state.p60z = '-'
    st.session_state.p60a = '-'
    st.session_state.p61 = '-'
    st.session_state.p61z = '-'
    st.session_state.p61a = '-'
    st.session_state.p62 = '-'
    st.session_state.p63 = '-'
    st.session_state.p63z = '-'
    st.session_state.p63a = '-'
    st.session_state.p63b = '-'
    st.session_state.p64 = '-'


    if st.session_state.participant_declaration == participant_options[0]:   #Below Level 1
        st.session_state.p58 = 'X'
    elif st.session_state.participant_declaration == participant_options[1]: #Level 1
        st.session_state.p59 = 'X'
    elif st.session_state.participant_declaration == participant_options[2]: #Level 2
        st.session_state.p60, st.session_state.p60z = 'X', 'X'
    elif st.session_state.participant_declaration == participant_options[3]: #Full Level 2
        st.session_state.p60, st.session_state.p60a = 'X', 'X'
    elif st.session_state.participant_declaration == participant_options[4]: #Level 3
        st.session_state.p61, st.session_state.p61z = 'X', 'X'
    elif st.session_state.participant_declaration == participant_options[5]: #Full Level 3
        st.session_state.p61, st.session_state.p61a = 'X', 'X'
    elif st.session_state.participant_declaration == participant_options[6]: #Level 4
        st.session_state.p62 = 'X'
    elif st.session_state.participant_declaration == participant_options[7]: #Level 5
        st.session_state.p63, st.session_state.p63z = 'X', 'X' 
    elif st.session_state.participant_declaration == participant_options[8]: #Level 6
        st.session_state.p63, st.session_state.p63a = 'X', 'X'
    elif st.session_state.participant_declaration == participant_options[9]: #Level 7 and above
        st.session_state.p63, st.session_state.p63b = 'X', 'X'
    elif st.session_state.participant_declaration == participant_options[10]: #No Qualifications
        st.session_state.p64 = 'X'
    


    # st.subheader('Training Providers declaration')
    # training_provider_options = [
    #     'Below Level 1', 'Level 1', 'Level 2', 'Level 3', 'Below Level 4',
    #     'Level 5 and above', 'No Qualifications', 'No Personal Learning Record'
    # ]

    # training_provider_declaration = st.radio(
    #     'Please check the PLR and record information about prior attainment level to ensure correct recording of prior attainment, as well as ensuring no duplication of learning aims or units takes place.',
    #     training_provider_options)
    # p65 = '-'
    # p66 = '-'
    # p67 = '-'
    # p68 = '-'
    # p69 = '-'
    # p70 = '-'
    # p71 = '-'
    # p72 = '-'
    # justification='-'


    # if training_provider_declaration == training_provider_options[0]:
    #     p65 = 'X'
    # elif training_provider_declaration == training_provider_options[1]:
    #     p66 = 'X'
    # elif training_provider_declaration == training_provider_options[2]:
    #     p67 = 'X'
    # elif training_provider_declaration == training_provider_options[3]:
    #     p68 = 'X'
    # elif training_provider_declaration == training_provider_options[4]:
    #     p69 = 'X'
    # elif training_provider_declaration == training_provider_options[5]:
    #     p70 = 'X'
    # elif training_provider_declaration == training_provider_options[6]:
    #     p71 = 'X'
    # elif training_provider_declaration == training_provider_options[7]:
    #     p72 = 'X'

    # justification = st.text_area(
    #         'If there is a discrepancy between Participant self declaration and the PLR, please record justification for level to be reported'
    #     )

    # st.subheader('Does the participant have Basic Skills?')

    # english_options = ['none', 'Entry Level', 'Level 1', 'Level 2+']

    # english_skill = st.selectbox('English', english_options)

    # p74 = '-'
    # p75 = '-'
    # p76 = '-'
    # p77 = '-'

    # if english_skill == english_options[0]:
    #     p74 = 'X'
    # elif english_skill == english_options[1]:
    #     p75 = 'X'
    # elif english_skill == english_options[2]:
    #     p76 = 'X'
    # elif english_skill == english_options[3]:
    #     p77 = 'X'

    # maths_options = ['none', 'Entry Level', 'Level 1', 'Level 2+']

    # maths_skill = st.selectbox('Maths', maths_options)

    # p78 = '-'
    # p79 = '-'
    # p80 = '-'
    # p81 = '-'

    # if maths_skill == maths_options[0]:
    #     p78 = 'X'
    # elif maths_skill == maths_options[1]:
    #     p79 = 'X'
    # elif maths_skill == maths_options[2]:
    #     p80 = 'X'
    # elif maths_skill == maths_options[3]:
    #     p81 = 'X'

    # esol_options = ['none', 'Entry Level', 'Level 1', 'Level 2+']

    # esol_skill = st.selectbox('ESOL', esol_options)

    # p82 = '-'
    # p83 = '-'
    # p84 = '-'
    # p85 = '-'

    # if esol_skill == esol_options[0]:
    #     p82 = 'X'
    # elif esol_skill == esol_options[1]:
    #     p83 = 'X'
    # elif esol_skill == esol_options[2]:
    #     p84 = 'X'
    # elif esol_skill == esol_options[3]:
    #     p85 = 'X'

    # st.subheader('Basic Skills Initial Assessment')
    # st.text(
    #     "Initial Assessment Outcomes – record the levels achieved by the Participant"
    # )

    # maths_options = ['-', 'E1', 'E2', 'E3', '1', '2']

    # maths_level = st.selectbox('Maths Level', maths_options)

    # p86 = ''

    # if maths_level in maths_options[1:]:
    #     p86 = maths_level

    # english_options = ['-', 'E1', 'E2', 'E3', '1', '2']

    # english_level = st.selectbox('English Level', english_options)

    # p87 = ''

    # if english_level in english_options[1:]:
    #     p87 = english_level

    # st.subheader('Numeracy and Literacy Programmes')
    # completion_programmes = st.radio(
    #     'Will the Participant be completing relevant Numeracy and/or Literacy programmes within their learning plan?',
    #     ['Yes', 'No'])
    # p88 = '-'
    # p89 = '-'

    # if completion_programmes == 'Yes':
    #     p88 = 'Y'
    #     p89 = '-'
    # elif completion_programmes == 'No':
    #     p88 = '-'
    #     p89 = 'N'

    # st.subheader('Additional Learning Support')
    # additional_support = st.radio(
    #     'Does the Participant require additional learning and/or learner support?',
    #     ['Yes', 'No'])
    # p90 = '-'
    # p91 = '-'
    # support_details = '-'

    # if additional_support == 'Yes':
    #     p90 = 'Y'
    #     p91 = '-'
    #     support_details = st.text_area(
    #         'If answered \'Yes\' above, please detail how the participant will be supported'
    #     )
    # elif additional_support == 'No':
    #     p90 = '-'
    #     p91 = 'N'

    # Navigation buttons
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        st.session_state.step = 9
        st.experimental_rerun()

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 7  # Go back to the previous step (Section 6)
        st.experimental_rerun()


elif st.session_state.step == 9:
    st.title("> 8: Current Skills, Experience, and IAG")

    st.header('Current Skills, Experience, and IAG')

    st.subheader('Highest Level of Education at start')
    education_options = [
        'Choose an option',  # Placeholder option
        'ISCED 0 - Lacking Foundation skills (below Primary Education)',
        'ISCED 1 - Primary Education',
        'ISCED 2 - GCSE D-G or 3-1/BTEC Level 1/Functional Skills Level 1',
        'ISCED 3 - GCSE A-C or 9-4/AS or A Level/NVQ or BTEC Level 2 or 3',
        'ISCED 4 - N/A',
        'ISCED 5 to 8 - BTEC Level 5 or NVQ Level 4, Foundation Degree, BA, MA or equivalent'
    ]

    # Change from selectbox to multiselect
    st.session_state.selected_levels = st.selectbox(
        'Select the highest level of education at start',
        options=education_options,
        index=0  # Default to the placeholder option
    )

    # Mandatory field validation
    if st.session_state.selected_levels == 'Choose an option':
        st.warning("Please select a valid education level before proceeding.")
        st.stop()

    # Initialize marks
    st.session_state.p93, st.session_state.p94, st.session_state.p95, st.session_state.p96, st.session_state.p97, st.session_state.p98 = '-', '-', '-', '-', '-', '-'

    # Mark selected options
    if education_options[0] in st.session_state.selected_levels:
        st.session_state.p93 = 'X'
    if education_options[1] in st.session_state.selected_levels:
        st.session_state.p94 = 'X'
    if education_options[2] in st.session_state.selected_levels:
        st.session_state.p95 = 'X'
    if education_options[3] in st.session_state.selected_levels:
        st.session_state.p96 = 'X'
    if education_options[4] in st.session_state.selected_levels:
        st.session_state.p97 = 'X'
    if education_options[5] in st.session_state.selected_levels:
        st.session_state.p98 = 'X'

    st.header('Other Information')

    st.session_state.job_role_activities='No job.'
    st.session_state.current_job = st.radio(
    'Are you currently doing job?',
    ['No', 'Yes'])
    if st.session_state.current_job=='Yes':
        st.subheader('Current Job Role and Day to Day Activities')
        st.session_state.job_role_activities = st.text_area(
            'What is your current job role and what are your day to day activities?'
        )


    st.subheader('Career Aspirations')
    st.session_state.career_aspirations = st.text_area('What are your career aspirations? (Please provide details.)')

    st.session_state.training_qualifications_needed='    '
    # st.subheader('Training/Qualifications Needed')
    # training_qualifications_needed = st.text_area(
    #     'What training/qualifications do you need to progress further in your career? (Planned and future training)'
    # )

    st.session_state.barriers_to_achieving_aspirations='    '
    # st.subheader('Barriers to Achieving Career Aspirations')
    # barriers_to_achieving_aspirations = st.text_area(
    #     'What are the barriers to achieving your career aspirations and goals?'
    # )

    # mandatory_fields.extend([f'p{i}' for i in range(99, 103)])

    # st.subheader('Courses/Programs Available')
    # courses_programs_available = st.text_area(
    #     'What courses/programs/activity are available to you in order to meet your and your employer\'s needs?'
    # )

    # st.header('Induction Checklist')


    # funded_by_mayor_of_london = st.checkbox(
    #     'This programme is funded by the Mayor of London')
    # lls_completed = st.checkbox(
    #     'The London Learning Survey (LLS) has been completed and submitted')
    # equality_diversity_policy = st.checkbox(
    #     'Equality and Diversity Policy/Procedure and point of contact')
    # health_safety_policy = st.checkbox(
    #     'Health and Safety Policy/Procedure and point of contact')
    # safeguarding_policy = st.checkbox(
    #     'Safeguarding Policy/Procedure and point of contact')
    # prevent_policy = st.checkbox(
    #     'PREVENT and point of contact (including British Values)')
    # disciplinary_policy = st.checkbox(
    #     'Disciplinary, Appeal and Grievance Policy/Procedures')
    # plagiarism_policy = st.checkbox('Plagiarism, Cheating Policy/Procedure')
    # terms_conditions = st.checkbox(
    #     'Terms and Conditions of Learning and programme content & programme delivery'
    # )

    # Navigation buttons
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        if (st.session_state.career_aspirations):
            st.session_state.step = 10
            st.experimental_rerun()
        else:
            st.warning("Please fill in all fields before proceeding.")

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 8  # Go back to the previous step (Section 6)
        st.experimental_rerun()


elif st.session_state.step == 10:
    st.title("> 9: Privacy Notice Text")

    # Privacy Notice Text
    privacy_notice = """
    Privacy Notice

    This privacy notice is issued by the Education and Skills Funding Agency (ESFA) on behalf of the Secretary of State for the Department of Education (DfE) to inform learners about the Individualised Learner Record (ILR) and how their personal information is used in the ILR. Your personal information is used by the DfE to exercise our functions under article 6(1)(e) of the UK GDPR and to meet our statutory responsibilities, including under the Apprenticeships, Skills, Children and Learning Act 2009.

    The ILR collects data about learners and learning undertaken. Publicly funded colleges, training organisations, local authorities, and employers (FE providers) must collect and return the data to the ESFA each year under the terms of a funding agreement, contract or grant agreement. It helps ensure that public money distributed through the ESFA is being spent in line with government targets. It is also used for education, training, employment, and wellbeing purposes, including research. We retain ILR learner data for 3 years for operational purposes and 66 years for research purposes. For more information about the ILR and the data collected, please see the ILR specification at https://www.gov.uk/government/collections/individualised-learner-record-ilr

    ILR data is shared with third parties where it complies with DfE data sharing procedures and where the law allows it. The DfE and the English European Social Fund (ESF) Managing Authority (or agents acting on their behalf) may contact learners to carry out research and evaluation to inform the effectiveness of training. In these cases, it is part of our statutory duties and we do not need your consent.

    For more information about how your personal data is used and your individual rights, please see the DfE Roles and Responsibilities Personal Information Charter(https://www.gov.uk/government/organisations/department-for-education/about/personal-information-charter) and the ESFA Privacy Notice (https://www.gov.uk/government/publications/esfa-privacy-notice).

    If you would like to get in touch with us, you can contact the DfE in the following ways:
    - Using our online contact form at https://www.gov.uk/government/organisations/department-for-education/about/personal-information-charter.
    - By telephoning the DfE Helpline on 0370 000 2288 or in writing to - Data Protection Officer, Ministerial and Public Communications Division, Department for Education, Piccadilly Gate, Store Street, Manchester, M1 2WD.

    By completing the 'Learner Declaration'. This means that:

    - You understand this provision is delivered by Prevista Ltd or by the named subcontractor on page 1 on behalf Prevista Ltd (or sub-contractor where indicated).
    - You will be the provider know of any changes in your personal circumstances.
    - You fully agree that the provider can process information about you.
    - You understand that the form will be kept until 31st December 2030 at the latest.

    Prevista Ltd will:

    - Provide appropriate guidance and support to the Subcontractor to ensure that they deliver high-quality services.
    - Monitor and evaluate the performance of the Subcontractor regularly to ensure that they meet the agreed-upon standards.
    - Provide the necessary resources and information to the Subcontractor to enable them to carry out their work effectively.
    - Ensure that the Subcontractor complies with all relevant laws and regulations.

    The Subcontractor will:

    - Deliver the agreed-upon services to a high standard and in a timely manner.
    - Comply with all relevant laws and regulations, including health and safety requirements.
    - Provide regular progress reports and updates to Prevista to ensure that they are kept informed of the work being carried out.
    - Work collaboratively with Prevista to ensure that the needs of students and other stakeholders are met.
    """


    # Privacy and Data Protection Information Section
    st.header('Privacy and Data Protection Information')

    # Display the privacy notice text as plain text
    st.text(privacy_notice)


    # Contact preferences
    st.write("Choose Y or N for any of the following if you AGREE to be contacted about courses/learning opportunities")
    st.session_state.contact_surveys = st.radio("For surveys & research", options=["Y", "N"])
    st.session_state.contact_phone = st.radio("Phone", options=["Y", "N"])
    st.session_state.contact_email = st.radio("Email", options=["Y", "N"])
    st.session_state.contact_post = st.radio("Post", options=["Y", "N"])
    # Initialize variables for contact preferences
    st.session_state.contact_surveys_val, st.session_state.contact_phone_val, st.session_state.contact_email_val, st.session_state.contact_post_val = '', '', '', ''
    # Update the variables based on user selections
    st.session_state.contact_surveys_val = 'Y' if st.session_state.contact_surveys == "Y" else 'N'
    st.session_state.contact_phone_val = 'Y' if st.session_state.contact_phone == "Y" else 'N'
    st.session_state.contact_email_val = 'Y' if st.session_state.contact_email == "Y" else 'N'
    st.session_state.contact_post_val = 'Y' if st.session_state.contact_post == "Y" else 'N'

    # Navigation buttons
    next_clicked = st.button("Next")
    back_clicked = st.button("Back")

    # Handle Next button click
    if next_clicked:
        st.session_state.step = 11
        st.experimental_rerun()

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 9  # Go back to the previous step (Section 6)
        st.experimental_rerun()


elif st.session_state.step == 11:
    st.title("Declaration and Signature")
    
    # Display the checked items
    st.subheader("Checked Items:")
    checked_items = [label for label, is_checked in st.session_state.checkboxes.items() if is_checked]
    if checked_items:
        for item in checked_items:
            st.write(f" - {item}")
    else:
        st.write("No checkboxes selected.")

    # Display all uploaded files
    st.subheader("Uploaded Files:")
    if st.session_state.files:
        # Use a set to track unique file names and create a list of unique file entries
        st.session_state.unique_files = []
        unique_file_names = set()

        for file_entry in st.session_state.files:
            if isinstance(file_entry, tuple):
                label, uploaded_file = file_entry
                if uploaded_file.name not in unique_file_names:
                    st.session_state.unique_files.append((label, uploaded_file))
                    unique_file_names.add(uploaded_file.name)
            else:
                if file_entry.name not in unique_file_names:
                    st.session_state.unique_files.append(file_entry)
                    unique_file_names.add(file_entry.name)

        # Display the unique files
        for file_entry in st.session_state.unique_files:
            if isinstance(file_entry, tuple):
                label, uploaded_file = file_entry
                st.write(f"- {label}: {uploaded_file.name}")
            else:
                st.write(f"- {file_entry.name}")
    else:
        st.write("No files uploaded.")


    # Declaration text
    st.subheader("Declaration:")
    st.text(
        'We hereby confirm that I have read, understood and agree with the contents of this document and above privacy notice, and understand that the programme is funded by the Mayor of London.'
    )
    st.header('Participant Declaration')
    st.text_area(
        'Declaration',
        f'I certify that I have provided all of the necessary information to confirm my eligibility for the Provision. \nI also consent for the named Training Provider to collect further evidence, from a 3rd party Training Provider, to support a progression claim on my behalf (where applicable).'

    )
    st.text("Participant Signature:")
    st.session_state.participant_signature_1 = st_canvas(
        fill_color="rgba(255, 255, 255, 1)",
        stroke_width=5,
        stroke_color="rgb(0, 0, 0)",  # Black stroke color
        background_color="#ffffcc",  # background color
        width=400,
        height=150,
        drawing_mode="freedraw",
        key='p'
    )
    # Set today's date automatically and display it
    st.session_state.date_signed = date.today().strftime("%d-%m-%Y")
    st.write(f"Date: **{st.session_state.date_signed}**")

    st.header('Training Provider Declarations')
    st.text_area(
        'Declaration',
        'I certify that I have seen and verified the supporting evidence as indicated above, to confirm the Participant eligibility for ESF funding and this specific project.'
    )

    st.session_state.tp_name = st.text_input('Name')
    st.session_state.tp_position = st.text_input('Position')
    # Validation to check if fields are empty
    if not st.session_state.tp_name or not st.session_state.tp_position:
        st.warning("Please fill in both Name and Position before proceeding.")
        is_button_disabled = True
    else:
        is_button_disabled = False

    st.text("Training Provider Signature:")
    st.session_state.participant_signature_2 = st_canvas(
        fill_color="rgba(255, 255, 255, 1)",
        stroke_width=5,
        stroke_color="rgb(0, 0, 0)",  # Black stroke color
        background_color="#ffffcc",  # background color
        width=400,
        height=150,
        drawing_mode="freedraw",
        key='tp'
    )

    # Set today's date automatically and display it
    st.write(f"Date: **{st.session_state.date_signed}**")


    # Submit button
    submit_clicked = st.button("Submit", disabled=is_button_disabled)    

# ####################################################################################################################################

    # Handle Submit button click
    if submit_clicked:
        st.warning('Please wait! We are currently processing. . . .', icon="🚨")
        
        try: # A joke
            response = requests.get("https://official-joke-api.appspot.com/random_joke")
            joke_data = response.json()
            setup = joke_data['setup']
            punchline = joke_data['punchline']
            
            st.write("A Joke:", setup)
            time.sleep(2)
            st.write('Punchline: ', punchline)
        except:
            st.warning("Unable to access Joke because of internet!")
        time.sleep(1)
        st.text('Processing . . . . . . . ')

    # if submit_button:
        st.session_state.placeholder_values = {

            'p241': st.session_state.learner_name,
            'p242': st.session_state.qualification,
            'p243': st.session_state.start_date,
            'p244': st.session_state.end_date,

            'p110': st.session_state.title_mr,
            'p111': st.session_state.title_mrs,
            'p112': st.session_state.title_miss,
            'p113': st.session_state.title_ms,

            'p1': st.session_state.first_name,
            'p2': st.session_state.middle_name,
            'p3': st.session_state.family_name,

            'p114': st.session_state.gender_m,
            'p115': st.session_state.gender_f,
            'p116': st.session_state.other_gender,
            'p117': st.session_state.other_gender_text,

            'p4': st.session_state.date_of_birth,

            'p118': st.session_state.current_age,
            'p119': st.session_state.ethnicity_vars['ethnicity_31'],
            'p120': st.session_state.ethnicity_vars['ethnicity_32'],
            'p121': st.session_state.ethnicity_vars['ethnicity_33'],
            'p122': st.session_state.ethnicity_vars['ethnicity_34'],
            'p123': st.session_state.ethnicity_vars['ethnicity_35'],
            'p124': st.session_state.ethnicity_vars['ethnicity_36'],
            'p125': st.session_state.ethnicity_vars['ethnicity_37'],
            'p126': st.session_state.ethnicity_vars['ethnicity_38'],
            'p127': st.session_state.ethnicity_vars['ethnicity_39'],
            'p128': st.session_state.ethnicity_vars['ethnicity_40'],
            'p129': st.session_state.ethnicity_vars['ethnicity_41'],
            'p130': st.session_state.ethnicity_vars['ethnicity_42'],
            'p131': st.session_state.ethnicity_vars['ethnicity_43'],
            'p132': st.session_state.ethnicity_vars['ethnicity_44'],
            'p133': st.session_state.ethnicity_vars['ethnicity_45'],
            'p134': st.session_state.ethnicity_vars['ethnicity_46'],
            'p135': st.session_state.ethnicity_vars['ethnicity_47'],
            'p136': st.session_state.ethnicity_vars['ethnicity_48'],
            'p137': st.session_state.national_insurance_number,
            'p138': st.session_state.house_no_name_street,
            'p139': st.session_state.suburb_village,
            'p140': st.session_state.town_city,
            'p141': st.session_state.county,
            'p142': st.session_state.country_of_domicile,
            'p143': st.session_state.current_postcode,
            'p144': st.session_state.postcode_prior_enrollment,
            'p145': st.session_state.email_address,
            'p146': st.session_state.primary_telephone_number,
            'p147': st.session_state.secondary_telephone_number,
            'p148': st.session_state.next_of_kin,
            'p149': st.session_state.emergency_contact_phone_number,

            'p150': st.session_state.no_member_employed_with_children,
            'p151': st.session_state.no_member_employed_without_children,
            'p152': st.session_state.single_adult_household_with_children,
            'p153': st.session_state.unemployed_single_adult_household,
            'p154': st.session_state.none_of_the_above,  

            'p155': st.session_state.has_disability,
            'p156': st.session_state.no_disability,

            'p157a': st.session_state.vision_impairment_primary,
            'p157b': st.session_state.vision_impairment_secondary,
            'p157c': st.session_state.vision_impairment_tertiary,
            'p158a': st.session_state.hearing_impairment_primary,
            'p158b': st.session_state.hearing_impairment_secondary,
            'p158c': st.session_state.hearing_impairment_tertiary,
            'p159a': st.session_state.mobility_impairment_primary,
            'p159b': st.session_state.mobility_impairment_secondary,
            'p159c': st.session_state.mobility_impairment_tertiary,
            'p160a': st.session_state.complex_disabilities_primary,
            'p160b': st.session_state.complex_disabilities_secondary,
            'p160c': st.session_state.complex_disabilities_tertiary,
            'p161a': st.session_state.social_emotional_difficulties_primary,
            'p161b': st.session_state.social_emotional_difficulties_secondary,
            'p161c': st.session_state.social_emotional_difficulties_tertiary,
            'p162a': st.session_state.mental_health_difficulty_primary,
            'p162b': st.session_state.mental_health_difficulty_secondary,
            'p162c': st.session_state.mental_health_difficulty_tertiary,
            'p163a': st.session_state.moderate_learning_difficulty_primary,
            'p163b': st.session_state.moderate_learning_difficulty_secondary,
            'p163c': st.session_state.moderate_learning_difficulty_tertiary,
            'p164a': st.session_state.severe_learning_difficulty_primary,
            'p164b': st.session_state.severe_learning_difficulty_secondary,
            'p164c': st.session_state.severe_learning_difficulty_tertiary,
            'p165a': st.session_state.dyslexia_primary,
            'p165b': st.session_state.dyslexia_secondary,
            'p165c': st.session_state.dyslexia_tertiary,
            'p166a': st.session_state.dyscalculia_primary,
            'p166b': st.session_state.dyscalculia_secondary,
            'p166c': st.session_state.dyscalculia_tertiary,
            'p167a': st.session_state.autism_spectrum_primary,
            'p167b': st.session_state.autism_spectrum_secondary,
            'p167c': st.session_state.autism_spectrum_tertiary,
            'p168a': st.session_state.aspergers_primary,
            'p168b': st.session_state.aspergers_secondary,
            'p168c': st.session_state.aspergers_tertiary,
            'p169a': st.session_state.temporary_disability_primary,
            'p169b': st.session_state.temporary_disability_secondary,
            'p169c': st.session_state.temporary_disability_tertiary,
            'p170a': st.session_state.speech_communication_needs_primary,
            'p170b': st.session_state.speech_communication_needs_secondary,
            'p170c': st.session_state.speech_communication_needs_tertiary,
            'p171a': st.session_state.physical_disability_primary,
            'p171b': st.session_state.physical_disability_secondary,
            'p171c': st.session_state.physical_disability_tertiary,
            'p172a': st.session_state.specific_learning_difficulty_primary,
            'p172b': st.session_state.specific_learning_difficulty_secondary,
            'p172c': st.session_state.specific_learning_difficulty_tertiary,
            'p173a': st.session_state.medical_condition_primary,
            'p173b': st.session_state.medical_condition_secondary,
            'p173c': st.session_state.medical_condition_tertiary,
            'p174a': st.session_state.other_learning_difficulty_primary,
            'p174b': st.session_state.other_learning_difficulty_secondary,
            'p174c': st.session_state.other_learning_difficulty_tertiary,
            'p175a': st.session_state.other_disability_primary,
            'p175b': st.session_state.other_disability_secondary,
            'p175c': st.session_state.other_disability_tertiary,
            'p176': st.session_state.prefer_not_to_say,
            'p177': st.session_state.additional_info,
            'p178': st.session_state.ex_offender_y,
            'p179': st.session_state.ex_offender_n,
            'p180': st.session_state.ex_offender_choose_not_to_say,

            'p189': st.session_state.homeless_y, 
            'p190': st.session_state.homeless_n,
            'p191': st.session_state.homeless_choose_not_to_say,

            'p181': st.session_state.internally_sourced_val,
            'p182': st.session_state.recommendation_val,
            'p183': st.session_state.event_val,
            'p184': st.session_state.self_referral_val,
            'p185': st.session_state.family_friends_val,
            'p186': st.session_state.other_val,
            'p187': st.session_state.website_val,
            'p188': st.session_state.promotional_material_val,
            'p188a': st.session_state.jobcentre_plus_val,

            'p192': st.session_state.unemployed_val,
            'p193': st.session_state.economically_inactive_val,
            'p194': st.session_state.employed_val,
            'p195': st.session_state.up_to_12_months_val,
            'p196': st.session_state.twelve_months_or_longer_val,
            'p197': st.session_state.jcp_dwp_val,
            'p198': st.session_state.careers_service_val,
            'p199': st.session_state.third_party_val,
            'p200': st.session_state.other_evidence_val,
            'p201': st.session_state.inactive_status_val,
            'p202': st.session_state.inactive_evidence_type_val,
            'p203': st.session_state.inactive_evidence_date_val,  
            'p204': st.session_state.employer_name_val,
            'p205': st.session_state.employer_address_1_val,
            'p206': st.session_state.employer_address_2_val,
            'p207': st.session_state.employer_address_3_val,
            'p208': st.session_state.employer_postcode_val,
            'p209': st.session_state.employer_contact_name_val,
            'p210': st.session_state.employer_contact_position_val,
            'p211': st.session_state.employer_contact_email_val,
            'p212': st.session_state.employer_contact_phone_val,
            'p213': st.session_state.employer_edrs_number_val,
            'p214': st.session_state.living_wage_val,
            'p215a': st.session_state.employment_hours_val_0,
            'p215b': st.session_state.employment_hours_val_6,
            'p216': st.session_state.claiming_benefits_val,
            'p217': st.session_state.sole_claimant_val,
            'p218': st.session_state.universal_credit_val,
            'p219': st.session_state.job_seekers_allowance_val,
            'p220': st.session_state.employment_support_allowance_val,
            'p221': st.session_state.incapacity_benefit_val,
            'p222': st.session_state.personal_independence_payment_val,
            'p223': st.session_state.other_benefit_val,
            'p224': st.session_state.benefit_claim_date_val,
            'p225': st.session_state.contact_surveys_val,
            'p226': st.session_state.contact_phone_val,
            'p227': st.session_state.contact_email_val,
            'p228': st.session_state.contact_post_val,

            'p5': st.session_state.nationality,
            'p6': st.session_state.full_uk_passport,
            'p7': st.session_state.full_eu_passport,
            'p8': st.session_state.national_identity_card,
            'p9': st.session_state.hold_settled_status,
            'p10': st.session_state.hold_pre_settled_status,
            'p11': st.session_state.hold_leave_to_remain,
            'p12': st.session_state.not_nationality,
            'p13': st.session_state.passport_non_eu,
            'p14': st.session_state.letter_uk_immigration,
            'p15': st.session_state.passport_endorsed,
            'p16': st.session_state.identity_card,
            'p17': st.session_state.country_of_issue,
            'p18': st.session_state.id_document_reference_number,
            'p19': st.session_state.e01_date_of_issue,
            'p20': st.session_state.e01_date_of_expiry,
            'p21': st.session_state.e01_additional_notes,
            'p22': st.session_state.full_passport_eu,
            'p23': st.session_state.national_id_card_eu,
            'p24': st.session_state.firearms_certificate,
            'p25': st.session_state.birth_adoption_certificate,
            'p26': st.session_state.e02_drivers_license,
            'p27': st.session_state.edu_institution_letter,
            'p28': st.session_state.e02_employment_contract,
            'p29': st.session_state.state_benefits_letter,
            'p30': st.session_state.pension_statement,
            'p31': st.session_state.northern_ireland_voters_card,
            'p32': st.session_state.e02_other_evidence_text,
            'p33': st.session_state.e02_date_of_issue,
            'p34': st.session_state.e03_drivers_license,
            'p35': st.session_state.bank_statement,
            'p36': st.session_state.pension_statement,
            'p37': st.session_state.mortgage_statement,
            'p38': st.session_state.utility_bill,
            'p39': st.session_state.council_tax_statement,
            'p40': st.session_state.electoral_role_evidence,
            'p41': st.session_state.homeowner_letter,
            'p42': st.session_state.e03_date_of_issue,
            'p43': st.session_state.e03_other_evidence_text,
            'p44': st.session_state.latest_payslip,
            'p45': st.session_state.e04_employment_contract,
            'p46': st.session_state.confirmation_from_employer,
            'p47': st.session_state.redundancy_notice,
            'p48': st.session_state.sa302_declaration,
            'p49': st.session_state.ni_contributions,
            'p50': st.session_state.business_records,
            'p51': st.session_state.companies_house_records,
            'p52': st.session_state.other_evidence_employed,
            'p53': st.session_state.unemployed,
            'p54': st.session_state.e04_date_of_issue,
            'p55': st.session_state.qualification_or_training_y,
            'p56': st.session_state.qualification_or_training_n,
            'p57': st.session_state.course_details + ' ' + st.session_state.funding_details,
            'p58': st.session_state.p58,
            'p59': st.session_state.p59,
            'p60': st.session_state.p60,
            'p61': st.session_state.p61,
            'p62': st.session_state.p62,
            'p63': st.session_state.p63,
            'p64': st.session_state.p64,

            'p60z' : st.session_state.p60z,
            'p60a' : st.session_state.p60a,
            'p61z' : st.session_state.p61z,
            'p61a' : st.session_state.p61a,
            'p63z' : st.session_state.p63z,
            'p63a' : st.session_state.p63a,
            'p63b' : st.session_state.p63b,


            'p65': st.session_state.selected_option,
            # 'p66': p66,
            # 'p67': p67,
            # 'p68': p68,
            # 'p69': p69,
            # 'p70': p70,
            # 'p71': p71,
            # 'p72': p72,
            # 'p73': justification,
            # 'p74': p74,
            # 'p75': p75,
            # 'p76': p76,
            # 'p77': p77,
            # 'p78': p78,
            # 'p79': p79,
            # 'p80': p80,
            # 'p81': p81,
            # 'p82': p82,
            # 'p83': p83,
            # 'p84': p84,
            # 'p85': p85,
            # 'p86': p86,
            # 'p87': p87,
            # 'p88': p88,
            # 'p89': p89,
            # 'p90': p90,
            # 'p91': p91,
            # 'p92': support_details,
            'p93': st.session_state.p93,
            'p94': st.session_state.p94,
            'p95': st.session_state.p95,
            'p96': st.session_state.p96,
            'p97': st.session_state.p97,
            'p98': st.session_state.p98,
            'p99': st.session_state.job_role_activities,
            'p100': st.session_state.career_aspirations,
            'p101': st.session_state.training_qualifications_needed,
            'p102': st.session_state.barriers_to_achieving_aspirations,
            # 'p103': courses_programs_available,
            # 'p113': participant_signature,
            'p231': st.session_state.date_signed,
            
            # for validation
            'p300': st.session_state.household_filled,
            'p301': st.session_state.e02_filled,
            'p302': st.session_state.e03_filled,
            'p303': len(st.session_state.selected_levels),
            # 'p304': referrall,
            'p305': st.session_state.specify_referral,
            'p232': st.session_state.tp_name,
            'p233': st.session_state.tp_position,

            'p235': st.session_state.job_position,
            'p236': st.session_state.job_start_date,
            
            'p237y': st.session_state.resident_y,
            'p237n': st.session_state.resident_n,
            'p238': st.session_state.country_of_birth,
            'p239': st.session_state.years_in_uk,

        }
        
        # Remove leading/trailing spaces, then replace internal spaces with underscores, and convert to lowercase
        safe_first_name = st.session_state.first_name.strip().replace(" ", "_").lower()
        safe_family_name = st.session_state.family_name.strip().replace(" ", "_").lower()

        # Define input and output paths
        template_file = "ph_esfa_v5.docx"
        modified_file = f"ESFA_Form_Submission_{sanitize_filename(safe_first_name)}_{sanitize_filename(safe_family_name)}.docx"

        # Define paths for both signatures
        signature_path_1 = fr'signature_1_{sanitize_filename(safe_first_name)}_{sanitize_filename(safe_family_name)}.png'
        resized_image_path_1 = fr'resized_signature_image_1_{sanitize_filename(safe_first_name)}_{sanitize_filename(safe_family_name)}.png'

        signature_path_2 = fr'signature_2_{sanitize_filename(safe_first_name)}_{sanitize_filename(safe_family_name)}.png'
        resized_image_path_2 = fr'resized_signature_image_2_{sanitize_filename(safe_first_name)}_{sanitize_filename(safe_family_name)}.png'

        # Check if the first signature exists in the session state
        if 'participant_signature_1' in st.session_state and len(st.session_state.participant_signature_1.json_data['objects']) != 0:
            try:
                # Convert the first drawing to a PIL image and save it
                signature_image_1 = PILImage.fromarray(
                    st.session_state.participant_signature_1.image_data.astype('uint8'), 'RGBA')
                signature_image_1.save(signature_path_1)

                # Open and resize the first image
                print(f"Opening image file: {signature_path_1}")
                resized_image_1 = PILImage.open(signature_path_1)
                print(f"Original image size (signature 1): {resized_image_1.size}")
                resized_image_1 = resize_image_to_fit_cell(resized_image_1, 200, 47)
                resized_image_1.save(resized_image_path_1)  # Save resized image to a file
                print(f"Resized image saved to: {resized_image_path_1}")
            except Exception as e:
                print(f"An error occurred while processing the first signature image: {e}")
                # Display the error message on the screen
                st.error('Please wait, form will reprocess and will give you the option again to submit in 10 SECONDS automatically')
                st.error(f"Please take screenshot of the following error and share with Developer: \n{str(e)}")
                time.sleep(12)

                st.session_state.submission_done = False
                st.session_state.step = 11
                st.experimental_rerun()

        else:
            st.warning("Participant's SIGNATURE is missing! Please draw the signature.")
            st.stop()

        # Check if the second signature exists in the session state
        if 'participant_signature_2' in st.session_state and len(st.session_state.participant_signature_2.json_data['objects']) != 0:
            try:
                # Convert the second drawing to a PIL image and save it
                signature_image_2 = PILImage.fromarray(
                    st.session_state.participant_signature_2.image_data.astype('uint8'), 'RGBA')
                signature_image_2.save(signature_path_2)

                # Open and resize the second image
                print(f"Opening image file: {signature_path_2}")
                resized_image_2 = PILImage.open(signature_path_2)
                print(f"Original image size (signature 2): {resized_image_2.size}")
                resized_image_2 = resize_image_to_fit_cell(resized_image_2, 200, 50)
                resized_image_2.save(resized_image_path_2)  # Save resized image to a file
                print(f"Resized image saved to: {resized_image_path_2}")
            except Exception as e:
                print(f"An error occurred while processing the first signature image: {e}")
                # Display the error message on the screen
                st.error('Please wait, form will reprocess and will give you the option again to submit in 10 SECONDS automatically')
                st.error(f"Please take screenshot of the following error and share with Developer: \n{str(e)}")
                time.sleep(12)

                st.session_state.submission_done = False
                st.session_state.step = 11
                st.experimental_rerun()
        else:
            st.warning("Training Provider's SIGNATURE is missing! Please draw the signature.")
            st.stop()
        
        # Call the function to replace placeholders with both resized images
        replace_placeholders(template_file, modified_file, st.session_state.placeholder_values, resized_image_path_1, resized_image_path_2)

        # Email

        # Sender email credentials

        # Credentials: Streamlit host st.secrets
        # sender_email = st.secrets["sender_email"]
        # sender_password = st.secrets["sender_password"]
        sender_email = get_secret("sender_email")
        sender_password = get_secret("sender_password")
        # sender_email = 'dummy'
        # sender_password = 'dummy'            

        receiver_email = sender_email
        
        # Credentials: Local env
        # load_dotenv()                                     # uncomment import of this library!
        # sender_email = os.getenv('EMAIL')
        # sender_password = os.getenv('PASSWORD')
        
        subject = f"ESFA: {st.session_state.selected_option} {st.session_state.first_name} {st.session_state.family_name} {date.today()} {st.session_state.specify_referral}"

        # Generate summary for checked items
        checked_items = [label for label, is_checked in st.session_state.checkboxes.items() if is_checked]
        checked_summary = "<br>".join([f"- {item}" for item in checked_items]) if checked_items else "No checkboxes selected."

        # Generate summary for uploaded files, ensuring uniqueness
        uploaded_files = list(set(file.name for file in st.session_state.files if isinstance(file, st.runtime.uploaded_file_manager.UploadedFile)))
        files_summary = "<br>".join([f"- {file}" for file in uploaded_files]) if uploaded_files else "No files uploaded."

        # Construct the email body with formatted sections in HTML
        body = f'''
        <p>ESFA Form submitted. Please find attached files.</p>

        <p><strong>Checked Items:</strong><br>
        {checked_summary}</p>

        <p><strong>Uploaded Files:</strong><br>
        {files_summary}</p>

        <p>Thank you.</p>
        '''

        # Local file path
        local_file_path = modified_file

        # Send email with attachments
        if st.session_state.files or local_file_path:
            # Remove duplicates while preserving order, using file name and size as the criteria
            seen = set()
            unique_files = []
            
            for file in st.session_state.files:
                file_identifier = (file.name, file.size)  # Use file name and size as a unique identifier
                if file_identifier not in seen:
                    unique_files.append(file)
                    seen.add(file_identifier)
            
            st.session_state.files = unique_files  # Update with the filtered list
            try:
                send_email_with_attachments(sender_email, sender_password, receiver_email, subject, body, st.session_state.files, local_file_path)
            except Exception as e:
                st.error(f"Failed to send email: {e}")

                # Provide file download button as a fallback
                st.warning("Email couldn't be sent, but you can download the file directly.")
                if local_file_path:
                    with open(local_file_path, 'rb') as f:
                        file_contents = f.read()
                        st.download_button(
                            label="Download Your File",
                            data=file_contents,
                            file_name=local_file_path.split('/')[-1],
                            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        )
                st.warning('Please wait, form will reprocess and will give you the option again to submit in 10 SECONDS')
                time.sleep(12)

                st.session_state.submission_done = False
                st.session_state.step = 11
                st.experimental_rerun()
                                        
            st.success("Submission Finished!")
            st.session_state.submission_done = True

            
        if st.session_state.submission_done:
            try:
                # file download button
                with open(modified_file, 'rb') as f:
                    file_contents = f.read()
                    st.download_button(
                        label="Download Your Response",
                        data=file_contents,
                        file_name=modified_file,
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )

                # clear session state
                st.session_state.files = []
                st.session_state.clear()
                st.write("Please close the form.")
                st.snow()

            except Exception as e:
                st.write("Unable to download the file. Please whatsapp learner name to +447405327072 for verificatino of submission.")
                st.error('Please wait, form will reprocess and will give you the option again to submit in 10 SECONDS')
                time.sleep(12)

                st.session_state.submission_done = False
                st.session_state.step = 11
                st.experimental_rerun()



            
            # st.experimental_rerun()  # Rerun the app to reflect the reset state
    

#111111111111111111
    # Add a warning before the back button
    st.info("If you go back, you will have to re-sign the form.")

    # Navigation buttons
    back_clicked = st.button("Back", disabled=st.session_state.submission_done)

    # Handle Back button click
    if back_clicked:
        st.session_state.step = 10  # Go back to the previous step
        st.experimental_rerun()
#11111111111111111


#         if st.button("Next"):
#             if (st.session_state.first_name):
#                 st.session_state.step = 12
#                 st.experimental_rerun()
#             else:
#                 st.warning("Please fill in all fields before proceeding.")



# elif st.session_state.step == 12:
#     st.title("> 11: Thank you for completing the enrollment form!")
#     st.write("We will process your application and get in touch with you soon.")
#     st.text("Prevista! Where future begins.")

# streamlit run app.py --server.port 8507
# Dev : https://linkedin.com/in/osamatech786